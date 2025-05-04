"""
Financial Document Intelligence System - Agent Tools

This module defines tools for the agents to use.
"""

import os
import json
import logging
import asyncio
import time
from typing import Dict, Any, List, Optional, Callable, Union
from datetime import datetime
from pathlib import Path
import csv
from io import StringIO
import aiohttp
import re
import string

from agents import Agent, function_tool, TContext as Context, Runner
from pydantic import BaseModel, Field
from fin_assistant import config as config
from fin_assistant.utils import load_json, save_json, get_file_info, chunk_text, call_ollama

logger = logging.getLogger(__name__)

# Progress callback for real-time updates
progress_callback: Optional[Callable] = None

# Define a base model class that forbids extra fields
class StrictBaseModel(BaseModel):
    model_config = {
        "extra": "forbid"
    }

# Directory Tools schema classes
class ScanDirectoryInput(StrictBaseModel):
    directory_path: str
    recursive: bool

class FileInfo(StrictBaseModel):
    name: str
    path: str
    size: int
    created: str
    extension: str

class ScanDirectoryOutput(StrictBaseModel):
    status: str
    directory: Optional[str] = None
    files_found: List[FileInfo] = []
    count: Optional[int] = None
    message: Optional[str] = None

# Directory Tools
@function_tool
async def scan_directory(params: ScanDirectoryInput) -> ScanDirectoryOutput:
    """
    Scan a directory for financial documents
    
    Args:
        params: Input parameters with directory_path and recursive flag
    
    Returns:
        Scan results including files found
    """
    logger.info(f"Scanning directory: {params.directory_path} (recursive={params.recursive})")
    
    files_found = []
    total_files = 0
    processed_files = 0
    
    try:
        # Handle case when directory doesn't exist
        if not os.path.exists(params.directory_path):
            if progress_callback:
                progress_callback({
                    'status': 'error',
                    'message': f"Directory not found: {params.directory_path}",
                    'progress': 0,
                    'total': 0,
                    'current_file': None
                })
            return ScanDirectoryOutput(
                status="error",
                message=f"Directory not found: {params.directory_path}",
                files_found=[]
            )
        
        # Count total files first for progress tracking
        if params.recursive:
            for root, _, files in os.walk(params.directory_path):
                for file in files:
                    extension = os.path.splitext(file)[1].lower()
                    if extension in config.ACCEPTED_EXTENSIONS:
                        total_files += 1
        else:
            for item in os.listdir(params.directory_path):
                file_path = os.path.join(params.directory_path, item)
                if os.path.isfile(file_path):
                    extension = os.path.splitext(item)[1].lower()
                    if extension in config.ACCEPTED_EXTENSIONS:
                        total_files += 1
        
        # If no files found, notify and return early
        if total_files == 0:
            if progress_callback:
                progress_callback({
                    'status': 'completed',
                    'message': f"No financial documents found in {params.directory_path}",
                    'progress': 0,
                    'total': 0,
                    'current_file': None
                })
            return ScanDirectoryOutput(
                status="success",
                directory=params.directory_path,
                files_found=[],
                count=0
            )
        
        # Immediate notification at start with total
        if progress_callback:
            progress_callback({
                'status': 'started',
                'message': f"Starting scan of {total_files} financial documents",
                'progress': 0,
                'total': total_files,
                'current_file': None
            })
            # Small delay to ensure the started message is processed
            await asyncio.sleep(0.5)
        
        # Walk through directory
        if params.recursive:
            for root, _, files in os.walk(params.directory_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    extension = os.path.splitext(file)[1].lower()
                    
                    if extension in config.ACCEPTED_EXTENSIONS:
                        file_info_dict = get_file_info(file_path)
                        file_info = FileInfo(
                            name=file_info_dict["name"],
                            path=file_info_dict["path"],
                            size=file_info_dict["size"],
                            created=file_info_dict["created"],
                            extension=file_info_dict["extension"]
                        )
                        files_found.append(file_info)
                        
                        processed_files += 1
                        
                        # Emit progress event with percentage
                        if progress_callback:
                            percent_complete = int((processed_files / total_files) * 100)
                            progress_callback({
                                'status': 'processing',
                                'message': f"Processing {percent_complete}% complete: {file}",
                                'current_file': file,
                                'progress': processed_files,
                                'total': total_files
                            })
                            # Delay between progress updates
                            await asyncio.sleep(0.2)
        else:
            for item in os.listdir(params.directory_path):
                file_path = os.path.join(params.directory_path, item)
                
                if os.path.isfile(file_path):
                    extension = os.path.splitext(item)[1].lower()
                    
                    if extension in config.ACCEPTED_EXTENSIONS:
                        file_info_dict = get_file_info(file_path)
                        file_info = FileInfo(
                            name=file_info_dict["name"],
                            path=file_info_dict["path"],
                            size=file_info_dict["size"],
                            created=file_info_dict["created"],
                            extension=file_info_dict["extension"]
                        )
                        files_found.append(file_info)
                        
                        processed_files += 1
                        
                        # Emit progress event with percentage
                        if progress_callback:
                            percent_complete = int((processed_files / total_files) * 100)
                            progress_callback({
                                'status': 'processing',
                                'message': f"Processing {percent_complete}% complete: {item}",
                                'current_file': item,
                                'progress': processed_files,
                                'total': total_files
                            })
                            # Delay between progress updates
                            await asyncio.sleep(0.2)
        
        # Final update at completion
        if progress_callback:
            # Small delay before final update
            await asyncio.sleep(0.3)
            progress_callback({
                'status': 'completed',
                'message': f"Completed scanning {processed_files} financial documents",
                'progress': processed_files,
                'total': total_files,
                'current_file': None
            })
            
        return ScanDirectoryOutput(
            status="success",
            directory=params.directory_path,
            files_found=files_found,
            count=len(files_found)
        )
    
    except Exception as e:
        logger.error(f"Error scanning directory {params.directory_path}: {str(e)}")
        if progress_callback:
            progress_callback({
                'status': 'error',
                'message': f"Error scanning directory: {str(e)}",
                'progress': processed_files,
                'total': total_files,
                'current_file': None
            })
        return ScanDirectoryOutput(
            status="error",
            message=str(e),
            files_found=files_found
        )

# List Files schema classes
class ListFilesInput(StrictBaseModel):
    directory_path: str
    file_extensions: Optional[List[str]] = None

class FileListItem(StrictBaseModel):
    name: str
    path: str
    extension: str
    size: int

class ListFilesOutput(StrictBaseModel):
    status: str
    directory: Optional[str] = None
    files: List[FileListItem] = []
    count: Optional[int] = None
    message: Optional[str] = None

@function_tool
async def list_files(params: ListFilesInput) -> ListFilesOutput:
    """
    List files in a directory with optional filtering by extension
    
    Args:
        params: Input parameters with directory_path and optional file_extensions
    
    Returns:
        List of files matching criteria
    """
    logger.info(f"Listing files in directory: {params.directory_path}")
    
    # Use default accepted extensions if none provided
    if params.file_extensions is None:
        file_extensions = config.ACCEPTED_EXTENSIONS
    else:
        file_extensions = params.file_extensions
    
    try:
        # Handle case when directory doesn't exist
        if not os.path.exists(params.directory_path):
            return ListFilesOutput(
                status="error",
                message=f"Directory not found: {params.directory_path}",
                files=[]
            )
        
        files = []
        for item in os.listdir(params.directory_path):
            file_path = os.path.join(params.directory_path, item)
            
            if os.path.isfile(file_path):
                extension = os.path.splitext(item)[1].lower()
                
                if extension in file_extensions:
                    files.append(FileListItem(
                        name=item,
                        path=file_path,
                        extension=extension,
                        size=os.path.getsize(file_path)
                    ))
        
        return ListFilesOutput(
            status="success",
            directory=params.directory_path,
            files=files,
            count=len(files)
        )
    
    except Exception as e:
        logger.error(f"Error listing files in {params.directory_path}: {str(e)}")
        return ListFilesOutput(
            status="error",
            message=str(e),
            files=[]
        )

# Document Tools schema classes
class ExtractTextInput(StrictBaseModel):
    file_path: str

class ExtractTextOutput(StrictBaseModel):
    status: str
    file_path: Optional[str] = None
    chunks: List[str] = []
    chunk_count: Optional[int] = None
    message: Optional[str] = None

class AnalyzeDocumentInput(StrictBaseModel):
    chunk: str

class AnalyzeDocumentOutput(StrictBaseModel):
    status: str
    analysis: str

# Document Tools
@function_tool 
async def extract_text(params: ExtractTextInput) -> ExtractTextOutput:
    """
    Extract text from a document file
    
    Args:
        params: Input parameters with file_path
    
    Returns:
        Extracted text chunks
    """
    logger.info(f"Extracting text from: {params.file_path}")
    
    try:
        # Check if file exists with more detailed error
        if not os.path.exists(params.file_path):
            logger.error(f"PDF file not found: {params.file_path}")
            return ExtractTextOutput(
                status="error",
                message=f"File not found: {params.file_path}. Please verify the file path is correct and the file exists.",
                chunks=[]
            )
        
        # Check if file is readable
        if not os.access(params.file_path, os.R_OK):
            logger.error(f"File not readable: {params.file_path}")
            return ExtractTextOutput(
                status="error",
                message=f"File exists but is not readable: {params.file_path}. Please check file permissions.",
                chunks=[]
            )
        
        file_extension = os.path.splitext(params.file_path)[1].lower()
        text = ""
        
        # Extract text based on file type
        if file_extension == ".pdf":
            # Use PyMuPDF (fitz)
            try:
                import fitz
                try:
                    doc = fitz.open(params.file_path)
                    for page_num in range(len(doc)):
                        page = doc.load_page(page_num)
                        text += page.get_text()
                    doc.close()
                except fitz.FileDataError as e:
                    logger.error(f"PDF file format error: {str(e)}")
                    return ExtractTextOutput(
                        status="error",
                        message=f"PDF file format error: {str(e)}. The file may be corrupted or not a valid PDF.",
                        chunks=[]
                    )
                except Exception as e:
                    logger.error(f"Error reading PDF: {str(e)}")
                    return ExtractTextOutput(
                        status="error",
                        message=f"Error reading PDF file: {str(e)}",
                        chunks=[]
                    )
            except ImportError:
                return ExtractTextOutput(
                    status="error",
                    message="PyMuPDF not installed. Install with: pip install PyMuPDF",
                    chunks=[]
                )
        
        elif file_extension == ".docx":
            # Use python-docx
            try:
                from docx import Document
                try:
                    doc = Document(params.file_path)
                    for para in doc.paragraphs:
                        text += para.text + "\n"
                except Exception as e:
                    logger.error(f"Error reading DOCX: {str(e)}")
                    return ExtractTextOutput(
                        status="error",
                        message=f"Error reading DOCX file: {str(e)}. The file may be corrupted.",
                        chunks=[]
                    )
            except ImportError:
                return ExtractTextOutput(
                    status="error",
                    message="python-docx not installed. Install with: pip install python-docx",
                    chunks=[]
                )
        
        elif file_extension in [".csv", ".xlsx"]:
            # Use pandas
            try:
                import pandas as pd
                try:
                    if file_extension == ".csv":
                        df = pd.read_csv(params.file_path, on_bad_lines='skip')
                    else:
                        df = pd.read_excel(params.file_path, engine='openpyxl')
                    text = df.to_string()
                except pd.errors.EmptyDataError:
                    logger.warning(f"CSV file is empty: {params.file_path}")
                    text = "Empty file"
                except Exception as e:
                    logger.error(f"Error reading {file_extension} file: {str(e)}")
                    return ExtractTextOutput(
                        status="error",
                        message=f"Error reading {file_extension} file: {str(e)}",
                        chunks=[]
                    )
            except ImportError:
                return ExtractTextOutput(
                    status="error",
                    message="pandas not installed. Install with: pip install pandas openpyxl",
                    chunks=[]
                )
        
        elif file_extension == ".json":
            # Read JSON file
            try:
                with open(params.file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    text = json.dumps(data, indent=2)
            except json.JSONDecodeError as e:
                logger.error(f"Invalid JSON file: {str(e)}")
                return ExtractTextOutput(
                    status="error",
                    message=f"Invalid JSON file: {str(e)}",
                    chunks=[]
                )
            except Exception as e:
                logger.error(f"Error reading JSON file: {str(e)}")
                return ExtractTextOutput(
                    status="error",
                    message=f"Error reading JSON file: {str(e)}",
                    chunks=[]
                )
        
        elif file_extension == ".txt":
            # Read text file
            try:
                with open(params.file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            except UnicodeDecodeError:
                # Try with a different encoding if UTF-8 fails
                try:
                    with open(params.file_path, 'r', encoding='latin-1') as f:
                        text = f.read()
                except Exception as e:
                    logger.error(f"Error reading text file with alternative encoding: {str(e)}")
                    return ExtractTextOutput(
                        status="error",
                        message=f"Error reading text file with alternative encoding: {str(e)}",
                        chunks=[]
                    )
            except Exception as e:
                logger.error(f"Error reading text file: {str(e)}")
                return ExtractTextOutput(
                    status="error",
                    message=f"Error reading text file: {str(e)}",
                    chunks=[]
                )
        
        else:
            return ExtractTextOutput(
                status="error",
                message=f"Unsupported file extension: {file_extension}",
                chunks=[]
            )
        
        # Check if we actually got any text
        if not text.strip():
            logger.warning(f"No text extracted from {params.file_path}")
            return ExtractTextOutput(
                status="warning",
                file_path=params.file_path,
                message="No text content extracted from file. The file may be empty, corrupt, or contain only images.",
                chunks=[]
            )
        
        # Chunk the text
        chunks = chunk_text(text, config.MAX_TOKENS, config.CHUNK_OVERLAP)
        
        return ExtractTextOutput(
            status="success",
            file_path=params.file_path,
            chunks=chunks,
            chunk_count=len(chunks)
        )
    
    except Exception as e:
        logger.error(f"Error extracting text from {params.file_path}: {str(e)}")
        return ExtractTextOutput(
            status="error",
            message=f"Error extracting text: {str(e)}",
            chunks=[]
        )

@function_tool
async def analyze_document_chunk(params: AnalyzeDocumentInput) -> AnalyzeDocumentOutput:
    """
    Analyze a document chunk
    
    Args:
        params: Input parameters with document chunk
    
    Returns:
        Analysis result
    """
    logger.info(f"Analyzing document chunk: {len(params.chunk)} characters")
    
    try:
        prompt = f"""Analyze the following financial document excerpt and extract key information in JSON format.
Focus on financial metrics, entities, dates, and important information.

DOCUMENT TEXT:
{params.chunk}

Provide the analysis in the following JSON format:
{{
  "report_type": "Income Statement/Balance Sheet/Cash Flow/etc.",
  "report_period": "Q1 2024/FY 2023/etc.",
  "client_name": "Company Name",
  "entity": "Finance Department/Accounting/etc.",
  "description": "Brief description of what this document contains",
  "information_present": ["revenue", "expenses", "profit", "assets", ...]
}}

JSON ANALYSIS:"""
        
        try:
            # Get analysis from Ollama
            response = await call_ollama(prompt, config.OLLAMA_MODEL)
            
            # Clean up the response - extract only the JSON part
            if "{" in response:
                json_start = response.find("{")
                json_end = response.rfind("}") + 1
                if json_start >= 0 and json_end > json_start:
                    response = response[json_start:json_end]
            
            # Try to parse as JSON to validate
            try:
                json.loads(response)
            except json.JSONDecodeError:
                # If not valid JSON, wrap in quotes as a string
                response = f'{{ "analysis": {json.dumps(response)} }}'
            
            return AnalyzeDocumentOutput(
                status="success",
                analysis=response
            )
        except ValueError as e:
            # Ollama service error
            logger.error(f"Ollama service error when analyzing document: {str(e)}")
            return AnalyzeDocumentOutput(
                status="error",
                analysis=f'{{"error": "Ollama service error: {str(e)}"}}'
            )
        except ConnectionError as e:
            # Ollama connection error
            logger.error(f"Ollama connection error when analyzing document: {str(e)}")
            return AnalyzeDocumentOutput(
                status="error",
                analysis=f'{{"error": "Ollama connection error: {str(e)}"}}'
            )
    except Exception as e:
        logger.error(f"Error analyzing document chunk: {str(e)}")
        return AnalyzeDocumentOutput(
            status="error",
            analysis=f'{{"error": "{str(e)}"}}'
        )

# CSV Tools schema classes
class Analysis(StrictBaseModel):
    analysis: str
    status: Optional[str] = None

class CSVGenerationInput(StrictBaseModel):
    file_info: FileInfo
    analysis: Analysis

class CSVGenerationOutput(StrictBaseModel):
    status: str
    csv_path: Optional[str] = None
    columns: Optional[List[str]] = None
    row_count: Optional[int] = None
    file_info: Optional[FileInfo] = None
    message: Optional[str] = None

# CSV Tools
@function_tool
async def generate_csv_from_analysis(params: CSVGenerationInput) -> CSVGenerationOutput:
    """
    Generate a CSV file from document analysis
    
    Args:
        params: Input parameters with file info and analysis
    
    Returns:
        Path to the generated CSV file
    """
    logger.info(f"Generating CSV from analysis for file: {params.file_info.name}")
    
    try:
        # Create file path for the CSV
        file_name = os.path.splitext(params.file_info.name)[0]
        csv_path = os.path.join(config.CSV_DIR, f"{file_name}_extracted.csv")
        
        # Get analysis and original file path
        analysis = params.analysis.analysis
        file_path = params.file_info.path
        
        # Extract text from the original file
        try:
            # Read the original file to get direct data for CSV
            file_ext = os.path.splitext(file_path)[1].lower()
            original_text = ""
            
            if file_ext == ".pdf":
                try:
                    import fitz
                    doc = fitz.open(file_path)
                    for page_num in range(len(doc)):
                        page = doc.load_page(page_num)
                        original_text += page.get_text()
                    doc.close()
                except Exception as e:
                    logger.warning(f"Could not extract PDF text: {str(e)}")
            elif file_ext == ".docx":
                try:
                    from docx import Document
                    doc = Document(file_path)
                    for para in doc.paragraphs:
                        original_text += para.text + "\n"
                except Exception as e:
                    logger.warning(f"Could not extract DOCX text: {str(e)}")
            elif file_ext == ".txt":
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        original_text = f.read()
                except Exception as e:
                    logger.warning(f"Could not read text file: {str(e)}")
            elif file_ext == ".csv":
                try:
                    import pandas as pd
                    df = pd.read_csv(file_path, on_bad_lines='skip')
                    original_text = df.to_string()
                except Exception as e:
                    logger.warning(f"Could not read CSV file: {str(e)}")
            elif file_ext == ".xlsx":
                try:
                    import pandas as pd
                    df = pd.read_excel(file_path, engine='openpyxl')
                    original_text = df.to_string()
                except Exception as e:
                    logger.warning(f"Could not read Excel file: {str(e)}")
            
            # Combine original text with analysis for better extraction
            combined_text = original_text + "\n\n" + analysis
        except Exception as e:
            logger.warning(f"Error reading original file: {str(e)}")
            combined_text = analysis
        
        # Try to determine document type from analysis to use appropriate CSV structure
        doc_type = "unknown"
        if "balance sheet" in analysis.lower() or "statement of financial position" in analysis.lower():
            doc_type = "balance_sheet"
        elif "income statement" in analysis.lower() or "profit and loss" in analysis.lower():
            doc_type = "income_statement"
        elif "cash flow" in analysis.lower() or "statement of cash flows" in analysis.lower():
            doc_type = "cash_flow"
        elif "annual report" in analysis.lower():
            doc_type = "annual_report"
        elif "bank statement" in analysis.lower() or "account statement" in analysis.lower():
            doc_type = "bank_statement"
        
        # Define standard columns based on document type (for error handling)
        if doc_type == "balance_sheet":
            standard_columns = ["Account", "Amount", "Category", "Type", "Notes"]
        elif doc_type == "income_statement":
            standard_columns = ["Item", "Amount", "Category", "Period", "Notes"]
        elif doc_type == "cash_flow":
            standard_columns = ["Activity", "Item", "Amount", "Category", "Notes"]
        elif doc_type == "bank_statement":
            standard_columns = ["Date", "Description", "Amount", "Type", "Balance"]
        elif doc_type == "annual_report":
            standard_columns = ["Section", "Item", "Amount", "Year", "Notes"]
        else:
            # Generic financial data format
            standard_columns = ["Category", "Item", "Amount", "Period", "Notes"]
        
        # ONLY APPROACH: Use the Ollama Agent extraction
        logger.info(f"Extracting data using Ollama agent for {file_name}")
        ollama_extraction_result = await extract_data_using_ollama_agent(combined_text, doc_type, params.file_info.name)
        
        if ollama_extraction_result and "data" in ollama_extraction_result and len(ollama_extraction_result["data"]) > 0:
            # Create CSV file using the Ollama-extracted data
            with open(csv_path, 'w', newline='') as csvfile:
                writer = csv.writer(csvfile)
                
                # Write header row - use columns from extraction result if available
                columns = ollama_extraction_result.get("columns", standard_columns)
                writer.writerow(columns)
                
                # Write data rows
                row_count = 0
                for row in ollama_extraction_result["data"]:
                    if isinstance(row, dict):
                        writer.writerow([row.get(col, "") for col in columns])
                        row_count += 1
                
                logger.info(f"Successfully generated CSV with {row_count} rows using Ollama agent extraction")
                
                # Always add document metadata as a separate section for reference
                writer.writerow([])  # Empty row as separator
                writer.writerow(["DOCUMENT METADATA", "", "", "", ""])
                writer.writerow(["File", params.file_info.name, "", "", ""])
                writer.writerow(["Path", params.file_info.path, "", "", ""])
                writer.writerow(["Size", str(params.file_info.size), "", "", ""])
                writer.writerow(["Created", params.file_info.created, "", "", ""])
                writer.writerow(["Document Type", doc_type.replace("_", " ").title(), "", "", ""])
                writer.writerow(["Extraction Time", datetime.now().isoformat(), "", "", ""])
            
            return CSVGenerationOutput(
                status="success",
                csv_path=csv_path,
                columns=columns,
                row_count=row_count,
                file_info=params.file_info
            )
        else:
            # If Ollama extraction failed, don't use fallbacks - just report the error
            logger.error(f"Ollama agent extraction failed for {file_name}. No fallback methods used.")
            
            # Create a simple CSV with error information
            try:
                # Ensure directory exists
                os.makedirs(os.path.dirname(csv_path), exist_ok=True)
                
                with open(csv_path, 'w', newline='') as csvfile:
                    writer = csv.writer(csvfile)
                    
                    # Header row
                    writer.writerow(standard_columns)
                    
                    # Error information
                    writer.writerow(["Error", "CSV Generation Failed", "Ollama extraction failed", "", "No fallback methods used"])
                    
                    # Add document metadata
                    writer.writerow([])  # Empty row as separator
                    writer.writerow(["DOCUMENT METADATA", "", "", "", ""])
                    writer.writerow(["File", params.file_info.name, "", "", ""])
                    writer.writerow(["Path", params.file_info.path, "", "", ""])
                    writer.writerow(["Size", str(params.file_info.size), "", "", ""])
                    writer.writerow(["Created", params.file_info.created, "", "", ""])
                    writer.writerow(["Document Type", doc_type.replace("_", " ").title(), "", "", ""])
                    writer.writerow(["Extraction Time", datetime.now().isoformat(), "", "", ""])
                
                return CSVGenerationOutput(
                    status="error",
                    csv_path=csv_path,
                    columns=standard_columns,
                    row_count=0,
                    file_info=params.file_info,
                    message="Ollama agent extraction failed. No fallback methods used."
                )
            except Exception as csv_error:
                logger.error(f"Failed to create error CSV: {str(csv_error)}")
                return CSVGenerationOutput(
                    status="error",
                    message=f"Failed to generate CSV: Ollama extraction failed. Additionally, could not create error CSV: {str(csv_error)}"
                )
    
    except Exception as e:
        logger.error(f"Error generating CSV: {str(e)}")
        return CSVGenerationOutput(
            status="error",
            message=f"Error generating CSV: {str(e)}"
        )

def extract_financial_data_by_pattern(writer, analysis, doc_type, columns):
    """Extract financial data from analysis text using pattern matching based on document type"""
    lines = analysis.split('\n')
    row_count = 0
    
    # Parse based on document type
    if doc_type == "balance_sheet":
        # Look for asset and liability patterns
        categories = {
            "asset": "Assets",
            "liability": "Liabilities", 
            "equity": "Equity",
            "current asset": "Assets",
            "fixed asset": "Assets",
            "current liability": "Liabilities",
            "long-term liability": "Liabilities"
        }
        
        for line in lines:
            # Try to match patterns like "Cash: $10,000" or "Total Assets: 50,000"
            if ':' in line:
                parts = line.split(':', 1)
                item = parts[0].strip()
                value_part = parts[1].strip()
                
                # Look for amount patterns in the value part
                amount = extract_amount(value_part)
                if amount:
                    # Determine category
                    category = "Unknown"
                    for key, val in categories.items():
                        if key in item.lower():
                            category = val
                            break
                    
                    # Determine type
                    type_val = ""
                    if "current" in item.lower():
                        type_val = "Current"
                    elif "long-term" in item.lower() or "fixed" in item.lower():
                        type_val = "Long-term"
                    
                    writer.writerow([item, amount, category, type_val, ""])
                    row_count += 1
    
    elif doc_type == "income_statement":
        # Look for revenue and expense patterns
        categories = {
            "revenue": "Revenue",
            "sales": "Revenue",
            "income": "Revenue",
            "expense": "Expense",
            "cost": "Expense",
            "profit": "Profit",
            "loss": "Loss",
            "margin": "Profit"
        }
        
        for line in lines:
            # Try to match patterns like "Revenue: $10,000" or "Net Income: 50,000"
            if ':' in line:
                parts = line.split(':', 1)
                item = parts[0].strip()
                value_part = parts[1].strip()
                
                # Look for amount patterns in the value part
                amount = extract_amount(value_part)
                if amount:
                    # Determine category
                    category = "Unknown"
                    for key, val in categories.items():
                        if key in item.lower():
                            category = val
                            break
                    
                    # Try to find period information (quarter/year)
                    period = ""
                    for period_pattern in ["Q1", "Q2", "Q3", "Q4", "FY", "2023", "2024"]:
                        if period_pattern in line:
                            period = period_pattern
                            # Try to find a year if we have a quarter
                            if period.startswith("Q") and any(year in line for year in ["2023", "2024", "2025"]):
                                for year in ["2023", "2024", "2025"]:
                                    if year in line:
                                        period += f" {year}"
                                        break
                            break
                    
                    writer.writerow([item, amount, category, period, ""])
                    row_count += 1
    
    elif doc_type == "cash_flow":
        # Look for cash flow patterns
        activities = {
            "operating": "Operating",
            "operation": "Operating",
            "investing": "Investing",
            "investment": "Investing",
            "financing": "Financing",
            "finance": "Financing"
        }
        
        for line in lines:
            # Try to match patterns like "Operating Cash Flow: $10,000"
            if ':' in line:
                parts = line.split(':', 1)
                item = parts[0].strip()
                value_part = parts[1].strip()
                
                # Look for amount patterns in the value part
                amount = extract_amount(value_part)
                if amount:
                    # Determine activity type
                    activity = "Other"
                    for key, val in activities.items():
                        if key in item.lower() or key in line.lower():
                            activity = val
                            break
                    
                    # Determine category (inflow/outflow)
                    category = "Unknown"
                    if amount.startswith("-"):
                        category = "Outflow"
                    else:
                        category = "Inflow"
                    
                    writer.writerow([activity, item, amount, category, ""])
                    row_count += 1
    
    elif doc_type == "bank_statement":
        # Try to extract transaction data
        for line in lines:
            # Look for date patterns (YYYY-MM-DD or MM/DD/YYYY)
            date_match = re.search(r'\d{4}-\d{2}-\d{2}|\d{1,2}/\d{1,2}/\d{4}|\d{1,2}/\d{1,2}/\d{2}', line)
            if date_match:
                date = date_match.group(0)
                # Extract the rest of the line as description
                description = line.replace(date, "").strip()
                
                # Look for amount patterns
                amount = extract_amount(line)
                
                # Determine transaction type
                type_val = ""
                if amount:
                    if amount.startswith("-"):
                        type_val = "Debit"
                    else:
                        type_val = "Credit"
                
                # Extract balance if present
                balance = ""
                balance_match = re.search(r'balance[:\s]+([0-9,.]+)', line.lower())
                if balance_match:
                    balance = balance_match.group(1).strip()
                
                writer.writerow([date, description, amount, type_val, balance])
                row_count += 1
    
    else:
        # Generic financial data extraction
        financial_items = [
            'revenue', 'income', 'expense', 'profit', 'loss', 'asset', 
            'liability', 'balance', 'total', 'net', 'cash', 'account',
            'equity', 'debt', 'loan', 'credit', 'debit', 'payment',
            'tax', 'dividend', 'interest', 'earning', 'cost', 'sales'
        ]
        
        for line in lines:
            # Check if line contains any financial terms
            if any(term in line.lower() for term in financial_items):
                # Try to extract item and amount
                if ':' in line:
                    parts = line.split(':', 1)
                    item = parts[0].strip()
                    value_part = parts[1].strip()
                    
                    # Look for amount patterns in the value part
                    amount = extract_amount(value_part)
                    if amount:
                        # Determine category
                        category = determine_financial_category(item.lower())
                        
                        # Try to find date information
                        date = ""
                        date_match = re.search(r'\d{4}-\d{2}-\d{2}|\d{1,2}/\d{1,2}/\d{4}|\d{1,2}/\d{1,2}/\d{2}', line)
                        if date_match:
                            date = date_match.group(0)
                        
                        writer.writerow([item, amount, category, date, ""])
                        row_count += 1
                
                # Also look for standalone amount patterns
                elif '$' in line or '€' in line or '£' in line:
                    # Look for currency amount patterns
                    amount = extract_amount(line)
                    if amount:
                        # Try to determine item and category
                        words = line.split()
                        item = ""
                        for i in range(len(words)):
                            if '$' in words[i] or '€' in words[i] or '£' in words[i]:
                                item = ' '.join(words[max(0, i-3):i]).strip()
                                break
                        
                        if not item:
                            item = "Unlabeled Amount"
                        
                        category = determine_financial_category(line.lower())
                        
                        # Try to find date information
                        date = ""
                        date_match = re.search(r'\d{4}-\d{2}-\d{2}|\d{1,2}/\d{1,2}/\d{4}|\d{1,2}/\d{1,2}/\d{2}', line)
                        if date_match:
                            date = date_match.group(0)
                        
                        writer.writerow([item, amount, category, date, ""])
                        row_count += 1
    
    return row_count

def extract_amount(text):
    """Extract a numerical amount from text"""
    # Look for currency patterns like $1,234.56 or 1,234.56 or 1234.56
    amount_match = re.search(r'[-+]?[$€£]?\s*[0-9,]+(\.[0-9]+)?', text)
    if amount_match:
        amount = amount_match.group(0)
        # Remove currency symbols and commas
        amount = amount.replace('$', '').replace('€', '').replace('£', '').replace(',', '').strip()
        return amount
    return ""

def determine_financial_category(text):
    """Determine financial category from text"""
    if any(term in text for term in ['revenue', 'income', 'sale', 'earning']):
        return "Revenue"
    elif any(term in text for term in ['expense', 'cost', 'payment']):
        return "Expense"
    elif any(term in text for term in ['profit', 'margin', 'net']):
        return "Profit"
    elif any(term in text for term in ['asset', 'receivable']):
        return "Asset"
    elif any(term in text for term in ['liability', 'payable', 'debt']):
        return "Liability"
    elif any(term in text for term in ['equity', 'capital', 'share']):
        return "Equity"
    elif any(term in text for term in ['cash', 'balance']):
        return "Cash"
    elif any(term in text for term in ['tax']):
        return "Tax"
    else:
        return "Other"

# Metadata Tools schema classes
class CreateMetadataInput(StrictBaseModel):
    file_info: FileInfo
    analysis: Analysis
    csv_path: Optional[str] = None

class MetadataEntry(StrictBaseModel):
    id: str
    file_name: str
    file_path: str
    file_size: Optional[int] = None
    file_extension: Optional[str] = None
    created_at: Optional[str] = None
    modified_at: Optional[str] = None
    report_type: str = "Unknown"
    report_period: str = "Unknown"
    client_name: str = "Unknown"
    entity: str = "Unknown"
    account_name: Optional[str] = None
    wallet_id: Optional[str] = None
    description: str = "No description available"
    information_present: List[str] = []
    csv_path: Optional[str] = None

class CreateMetadataOutput(StrictBaseModel):
    status: str
    metadata: Optional[MetadataEntry] = None
    file_info: Optional[FileInfo] = None
    message: Optional[str] = None

# Metadata Tools
@function_tool
async def create_metadata(params: CreateMetadataInput) -> CreateMetadataOutput:
    """
    Create metadata for a document
    
    Args:
        params: Input parameters with file_info, analysis, and optional csv_path
    
    Returns:
        Metadata creation results
    """
    logger.info(f"Creating metadata for: {params.file_info.name}")
    
    try:
        # Extract analysis to structured metadata
        analysis_text = params.analysis.analysis
        
        # ONLY APPROACH: Use the Ollama Agent extraction for metadata
        logger.info(f"Extracting metadata using Ollama agent for {params.file_info.name}")
        extracted_metadata = await extract_metadata_using_ollama_agent(params.file_info, analysis_text)
        
        if not extracted_metadata:
            # If Ollama extraction failed, report the error - no fallbacks
            logger.error(f"Ollama agent extraction failed for metadata: {params.file_info.name}. No fallback methods used.")
            return CreateMetadataOutput(
                status="error",
                message="Ollama agent extraction failed for metadata. No fallback methods used.",
                file_info=params.file_info
            )
        
        # Generate CSV path based on file name
        csv_file_name = os.path.splitext(params.file_info.name)[0] + "_extracted.csv"
        csv_file_path = os.path.join(config.CSV_DIR, csv_file_name)
        
        # Create metadata entry with the extracted information
        from fin_assistant.utils import generate_id
        
        metadata_entry = MetadataEntry(
            id=generate_id(),
            file_name=params.file_info.name,
            file_path=params.file_info.path,
            file_size=params.file_info.size,
            file_extension=params.file_info.extension,
            created_at=params.file_info.created,
            modified_at=datetime.now().isoformat(),
            report_type=extracted_metadata.get("report_type", "Unknown"),
            report_period=extracted_metadata.get("report_period", "Unknown"),
            client_name=extracted_metadata.get("client_name", "Unknown"),
            entity=extracted_metadata.get("entity", "Unknown"),
            account_name=extracted_metadata.get("account_name", "Unknown"),
            wallet_id=extracted_metadata.get("wallet_id", f"WLT_{int(time.time()) % 10000:04d}"),
            description=extracted_metadata.get("description", "Unknown"),
            information_present=extracted_metadata.get("information_present", []),
            csv_path=params.csv_path or csv_file_path
        )
        
        # Load existing metadata
        metadata_json = load_json(config.METADATA_PATH, {"documents": []})
        
        # Check if document with same path already exists
        existing_index = None
        for i, doc in enumerate(metadata_json["documents"]):
            if doc.get("file_path") == params.file_info.path:
                existing_index = i
                break
        
        # Update or append metadata
        metadata_dict = metadata_entry.model_dump()
        if existing_index is not None:
            # Update existing entry
            metadata_json["documents"][existing_index] = metadata_dict
        else:
            # Add new entry
            metadata_json["documents"].append(metadata_dict)
        
        # Save metadata
        save_json(config.METADATA_PATH, metadata_json)
        
        return CreateMetadataOutput(
            status="success",
            metadata=metadata_entry,
            file_info=params.file_info
        )
    
    except Exception as e:
        logger.error(f"Error creating metadata: {str(e)}")
        return CreateMetadataOutput(
            status="error",
            message=str(e),
            file_info=params.file_info
        )

# Query Tools schema classes
class SearchDocumentsInput(StrictBaseModel):
    query: str
    limit: int = 5

class SearchResult(StrictBaseModel):
    id: str
    file_name: str
    file_path: str
    report_type: str = "Unknown"
    report_period: str = "Unknown"
    client_name: str = "Unknown"
    entity: str = "Unknown"
    description: str = "No description available"
    csv_path: Optional[str] = None
    score: int = 0

class SearchDocumentsOutput(StrictBaseModel):
    status: str
    query: str
    results: List[SearchResult] = []
    count: int = 0
    message: Optional[str] = None

class GetDocumentInput(StrictBaseModel):
    doc_id: str

class GetDocumentOutput(StrictBaseModel):
    status: str
    document: Optional[Dict[str, Any]] = None
    message: Optional[str] = None

class AnswerQueryInput(StrictBaseModel):
    doc_id: str
    query: str

class AnswerQueryOutput(StrictBaseModel):
    status: str
    query: str
    answer: Optional[str] = None
    document: Optional[Dict[str, Any]] = None
    csv_path: Optional[str] = None
    file_path: Optional[str] = None
    message: Optional[str] = None

# Query Tools
@function_tool
async def search_documents(params: SearchDocumentsInput) -> SearchDocumentsOutput:
    """
    Search for documents matching a query
    
    Args:
        params: Input parameters with query and limit
    
    Returns:
        Search results
    """
    logger.info(f"Searching documents for: {params.query}")
    
    try:
        # Load metadata
        metadata_json = load_json(config.METADATA_PATH, {"documents": []})
        documents = metadata_json.get("documents", [])
        
        if not documents:
            return SearchDocumentsOutput(
                status="success",
                query=params.query,
                message="No documents found in metadata",
                results=[]
            )
        
        # Normalize query
        query_lower = params.query.lower()
        
        # Calculate relevance scores
        results = []
        for doc in documents:
            score = 0
            
            # Check report type
            if "report_type" in doc and query_lower in doc["report_type"].lower():
                score += 5
            
            # Check client name
            if "client_name" in doc and query_lower in doc["client_name"].lower():
                score += 4
            
            # Check report period
            if "report_period" in doc and query_lower in doc["report_period"].lower():
                score += 4
            
            # Check entity
            if "entity" in doc and query_lower in doc["entity"].lower():
                score += 3
            
            # Check description
            if "description" in doc and query_lower in doc["description"].lower():
                score += 2
            
            # Check information present
            for info in doc.get("information_present", []):
                if query_lower in info.lower():
                    score += 2
            
            # Include document if score > 0
            if score > 0:
                results.append(SearchResult(
                    id=doc["id"],
                    file_name=doc["file_name"],
                    file_path=doc["file_path"],
                    report_type=doc.get("report_type", "Unknown"),
                    report_period=doc.get("report_period", "Unknown"),
                    client_name=doc.get("client_name", "Unknown"),
                    entity=doc.get("entity", "Unknown"),
                    description=doc.get("description", "No description available"),
                    csv_path=doc.get("csv_path"),
                    score=score
                ))
        
        # Sort by score (descending)
        results.sort(key=lambda x: x.score, reverse=True)
        
        # Limit results
        results = results[:params.limit]
        
        return SearchDocumentsOutput(
            status="success",
            query=params.query,
            results=results,
            count=len(results)
        )
    
    except Exception as e:
        logger.error(f"Error searching documents: {str(e)}")
        return SearchDocumentsOutput(
            status="error",
            query=params.query,
            message=str(e),
            results=[]
        )

@function_tool
async def get_document_by_id(params: GetDocumentInput) -> GetDocumentOutput:
    """
    Get a document by its ID
    
    Args:
        params: Input parameters with document ID
    
    Returns:
        Document metadata
    """
    logger.info(f"Getting document by ID: {params.doc_id}")
    
    try:
        # Load metadata
        metadata_json = load_json(config.METADATA_PATH, {"documents": []})
        
        # Find document by ID
        for doc in metadata_json.get("documents", []):
            if doc.get("id") == params.doc_id:
                return GetDocumentOutput(
                    status="success",
                    document=doc
                )
        
        return GetDocumentOutput(
            status="error",
            message=f"Document not found with ID: {params.doc_id}"
        )
    
    except Exception as e:
        logger.error(f"Error getting document by ID: {str(e)}")
        return GetDocumentOutput(
            status="error",
            message=str(e)
        )

@function_tool
async def answer_query(params: AnswerQueryInput) -> AnswerQueryOutput:
    """
    Answer a query using a document
    
    Args:
        params: Input parameters with document ID and query
    
    Returns:
        Answer to the query
    """
    logger.info(f"Answering query using document {params.doc_id}: {params.query}")
    
    try:
        # Get document metadata
        doc_result = await get_document_by_id(GetDocumentInput(doc_id=params.doc_id))
        
        if doc_result.status != "success":
            return AnswerQueryOutput(
                status="error",
                query=params.query,
                message=doc_result.message
            )
        
        document = doc_result.document
        
        # Check if CSV exists
        csv_path = document.get("csv_path")
        if not csv_path or not os.path.exists(csv_path):
            return AnswerQueryOutput(
                status="error",
                query=params.query,
                message=f"CSV file not found for document ID: {params.doc_id}"
            )
        
        # Read CSV
        import pandas as pd
        df = pd.read_csv(csv_path)
        
        # Convert to string for prompt
        csv_data = df.to_string(index=False)
        
        # Create prompt for Ollama
        prompt = f"""You are a financial expert. Given the below CSV table and a question, return the most accurate answer.

DOCUMENT INFO:
Type: {document.get('report_type', 'Unknown')}
Period: {document.get('report_period', 'Unknown')}
Client: {document.get('client_name', 'Unknown')}
Description: {document.get('description', 'No description available')}

CSV DATA:
{csv_data}

QUESTION:
{params.query}

ANSWER:
"""
        
        try:
            # Call Ollama
            response = await call_ollama(prompt, config.OLLAMA_MODEL)
            
            return AnswerQueryOutput(
                status="success",
                query=params.query,
                answer=response,
                document=document,
                csv_path=csv_path,
                file_path=document.get("file_path")
            )
        except ValueError as e:
            # Ollama service error
            logger.error(f"Ollama service error: {str(e)}")
            return AnswerQueryOutput(
                status="error",
                query=params.query,
                message=f"Ollama service error: {str(e)}"
            )
        except ConnectionError as e:
            # Ollama connection error
            logger.error(f"Ollama connection error: {str(e)}")
            return AnswerQueryOutput(
                status="error",
                query=params.query,
                message=f"Ollama connection error: {str(e)}"
            )
    
    except Exception as e:
        logger.error(f"Error answering query: {str(e)}")
        return AnswerQueryOutput(
            status="error",
            query=params.query,
            message=str(e)
        )

# Add Pydantic model for Ollama API call
class OllamaApiInput(StrictBaseModel):
    prompt: str
    model: str = "llama2"

class OllamaApiOutput(StrictBaseModel):
    status: str
    response: Optional[str] = None
    model: Optional[str] = None
    message: Optional[str] = None

# Ollama Tools
@function_tool
async def call_ollama_api(params: OllamaApiInput) -> OllamaApiOutput:
    """
    Call Ollama API directly
    
    Args:
        params: Input parameters with prompt and model name
    
    Returns:
        Ollama API response
    """
    logger.info(f"Calling Ollama API with model: {params.model}")
    
    try:
        response = await call_ollama(params.prompt, params.model)
        
        return OllamaApiOutput(
            status="success",
            response=response,
            model=params.model
        )
    
    except ValueError as e:
        # Ollama service error
        logger.error(f"Ollama service error: {str(e)}")
        return OllamaApiOutput(
            status="error",
            message=f"Ollama service error: {str(e)}"
        )
    except ConnectionError as e:
        # Ollama connection error
        logger.error(f"Ollama connection error: {str(e)}")
        return OllamaApiOutput(
            status="error",
            message=f"Ollama connection error: {str(e)}"
        )
    except Exception as e:
        logger.error(f"Error calling Ollama API: {str(e)}")
        return OllamaApiOutput(
            status="error",
            message=f"Unexpected error: {str(e)}"
        )

# Add these Pydantic models for the tools in registration functions
class ScanDirectoryToolInput(StrictBaseModel):
    directory_path: str
    recursive: bool = False

class AnalyzeDocumentToolInput(StrictBaseModel):
    file_path: str

class SimpleDataType(StrictBaseModel):
    value: str = ""
    description: str = ""
    metadata: Optional[str] = None

class GenerateCsvToolInput(StrictBaseModel):
    file_path: str
    data: SimpleDataType

class AskOllamaInput(StrictBaseModel):
    prompt: str
    model: Optional[str] = None

class CreateMetadataToolInput(StrictBaseModel):
    file_path: str
    document_type: str

class AnalyzeFileTypeInput(StrictBaseModel):
    file_path: str

class AnswerQueryToolInput(StrictBaseModel):
    query: str
    doc_id: Optional[str] = None

class GenericToolOutput(StrictBaseModel):
    status: str
    message: Optional[str] = None
    data: Optional[Dict[str, Any]] = None

# Replace the register_directory_tools function
def register_directory_tools(agent: Agent) -> None:
    """
    Register directory scanner tools for an agent
    
    Args:
        agent: Agent instance to register tools with
    """
    # Create tool functions
    @function_tool
    async def scan_directory_tool(params: ScanDirectoryToolInput) -> ScanDirectoryOutput:
        """
        Scan a directory for financial documents
        
        Args:
            params: Input parameters with directory path and recursive flag
            
        Returns:
            Scan results including file information
        """
        return await scan_directory(params)
    
    # Add the tool directly to the agent's tools list
    agent.tools.append(scan_directory_tool)
    
    logger.info(f"Registered directory tools for {agent.name}")

# Replace the register_document_tools function
def register_document_tools(agent: Agent) -> None:
    """
    Register document analyzer tools for an agent
    
    Args:
        agent: Agent instance to register tools with
    """
    # Create tool to analyze documents
    @function_tool
    async def analyze_document_tool(params: AnalyzeDocumentToolInput) -> GenericToolOutput:
        """
        Analyze a document's content to extract structured information
        
        Args:
            params: Input parameters with file path
            
        Returns:
            Document analysis results
        """
        result = await analyze_document(params.file_path)
        return GenericToolOutput(
            status="success",
            data=result
        )
    
    # Add the tool directly to the agent's tools list
    agent.tools.append(analyze_document_tool)
    
    logger.info(f"Registered document tools for {agent.name}")

# Replace the register_csv_tools function
def register_csv_tools(agent: Agent) -> None:
    """
    Register CSV tools for an agent
    
    Args:
        agent: Agent instance to register tools with
    """
    # Create tool to generate CSV
    @function_tool
    async def generate_csv_tool(params: GenerateCsvToolInput) -> GenericToolOutput:
        """
        Generate a CSV file from extracted document data
        
        Args:
            params: Input parameters with file path and extracted data
            
        Returns:
            CSV generation results
        """
        # Convert SimpleDataType to dict for backward compatibility
        data_dict = {
            "value": params.data.value,
            "description": params.data.description
        }
        if params.data.metadata:
            data_dict["metadata"] = params.data.metadata
            
        result = await generate_csv(params.file_path, data_dict)
        return GenericToolOutput(
            status="success",
            data=result
        )
    
    # Add the tool directly to the agent's tools list
    agent.tools.append(generate_csv_tool)
    
    logger.info(f"Registered CSV tools for {agent.name}")

# Replace the register_ollama_tools function
def register_ollama_tools(agent: Agent) -> None:
    """
    Register Ollama tools for an agent
    
    Args:
        agent: Agent instance to register tools with
    """
    # Create Ollama tool
    @function_tool
    async def ask_ollama(params: AskOllamaInput) -> GenericToolOutput:
        """
        Ask Ollama a question and get a response
        
        Args:
            params: Input parameters with prompt and optional model name
            
        Returns:
            Ollama's response
        """
        model_name = params.model or config.OLLAMA_MODEL
        response = await call_ollama(params.prompt, model_name)
        return GenericToolOutput(
            status="success",
            data={"response": response, "model": model_name}
        )
    
    # Add the tool directly to the agent's tools list
    agent.tools.append(ask_ollama)
    
    logger.info(f"Registered Ollama tools for {agent.name}")

# Replace the register_metadata_tools function
def register_metadata_tools(agent: Agent) -> None:
    """
    Register metadata tools for an agent
    
    Args:
        agent: Agent instance to register tools with
    """
    # Create metadata tool with simpler structure
    @function_tool
    async def create_metadata_tool(params: CreateMetadataToolInput) -> GenericToolOutput:
        """
        Create metadata for a document
        
        Args:
            params: Input parameters with file path and document type
            
        Returns:
            Metadata creation results
        """
        logger.info(f"Creating metadata for: {params.file_path}")
        try:
            # Get file info
            file_info = get_file_info(params.file_path)
            
            # Create simple metadata
            metadata = {
                "id": generate_id(),
                "file_name": file_info["name"],
                "file_path": file_info["path"],
                "file_size": file_info["size"],
                "file_extension": file_info["extension"],
                "created_at": file_info["created"],
                "modified_at": datetime.now().isoformat(),
                "report_type": params.document_type,
                "extraction_time": datetime.now().isoformat()
            }
            
            # Save metadata to file
            metadata_path = config.METADATA_PATH
            current_metadata = load_json(metadata_path, {"documents": []})
            
            # Add metadata to documents list
            if "documents" not in current_metadata:
                current_metadata["documents"] = []
                
            current_metadata["documents"].append(metadata)
            save_json(metadata_path, current_metadata)
            
            return GenericToolOutput(
                status="success",
                message=f"Metadata created for {params.file_path}",
                data={"metadata": metadata}
            )
        except Exception as e:
            logger.error(f"Error creating metadata: {str(e)}")
            return GenericToolOutput(
                status="error",
                message=f"Error creating metadata: {str(e)}"
            )
    
    # Add the tool directly to the agent's tools list
    agent.tools.append(create_metadata_tool)
    
    logger.info(f"Registered metadata tools for {agent.name}")

# Replace the register_file_analyzer_tools function
def register_file_analyzer_tools(agent: Agent) -> None:
    """
    Register file analyzer tools for an agent
    
    Args:
        agent: Agent instance to register tools with
    """
    # Create document type analyzer tool
    @function_tool
    async def analyze_file_type_tool(params: AnalyzeFileTypeInput) -> DocumentTypeAnalysisOutput:
        """
        Analyze a document to determine its report type
        
        Args:
            params: Input parameters with file path
            
        Returns:
            Document type analysis results
        """
        # Convert path to string if needed
        file_path = str(params.file_path) if not isinstance(params.file_path, str) else params.file_path
        
        # Call analyze_document_type and await the result
        return await analyze_document_type(file_path)
    
    # Add the tool directly to the agent's tools list
    agent.tools.append(analyze_file_type_tool)
    
    logger.info(f"Registered file analyzer tools for {agent.name}")

# Replace the register_query_tools function
def register_query_tools(agent: Agent) -> None:
    """
    Register query tools for an agent
    
    Args:
        agent: Agent instance to register tools with
    """
    # Create query tool
    @function_tool
    async def answer_query_tool(params: AnswerQueryToolInput) -> GenericToolOutput:
        """
        Answer a query about financial documents
        
        Args:
            params: Input parameters with query text and optional document ID
            
        Returns:
            Query answer results
        """
        # Pass the params object directly to answer_query
        result = await answer_query(params)
        return GenericToolOutput(
            status="success",
            data=result
        )
    
    # Add the tool directly to the agent's tools list
    agent.tools.append(answer_query_tool)
    
    logger.info(f"Registered query tools for {agent.name}")

# Helper function to generate unique IDs (missing from the code snippets)
def generate_id() -> str:
    """Generate a unique ID for metadata entries"""
    return f"doc_{int(time.time() * 1000)}"

# Add document type analysis model
class DocumentTypeAnalysisInput(StrictBaseModel):
    file_path: str
    content: Optional[str] = None

class DocumentTypeAnalysisOutput(StrictBaseModel):
    file_path: str
    document_type: str
    confidence: float
    detected_fields: List[str] = []
    metadata: Dict[str, Any] = {}

async def analyze_document_type(file_path: str, content: Optional[str] = None) -> DocumentTypeAnalysisOutput:
    """
    Analyze a document to determine its report type and extract comprehensive metadata
    
    Args:
        file_path: Path to the document
        content: Document content if already extracted
        
    Returns:
        Document type information and metadata for all detected reports
    """
    # Convert Path object to string if needed
    if not isinstance(file_path, str):
        file_path = str(file_path)
    
    try:
        # Check if we have a cached result for this file
        import hashlib
        file_hash = None
        try:
            with open(file_path, 'rb') as f:
                file_hash = hashlib.md5(f.read()).hexdigest()
            
            # Define cache path
            cache_dir = os.path.join(config.DATA_DIR, 'cache')
            os.makedirs(cache_dir, exist_ok=True)
            cache_file = os.path.join(cache_dir, f"{file_hash}.json")
            
            # Check if we have a cached analysis
            if os.path.exists(cache_file):
                logger.info(f"Using cached analysis for {file_path}")
                with open(cache_file, 'r') as f:
                    cached_data = json.load(f)
                    return DocumentTypeAnalysisOutput(
                        file_path=file_path,
                        document_type=cached_data.get("document_type", "unknown"),
                        confidence=cached_data.get("confidence", 0.5),
                        detected_fields=cached_data.get("detected_fields", []),
                        metadata=cached_data.get("metadata", {})
                    )
        except Exception as e:
            logger.warning(f"Couldn't check cache for {file_path}: {str(e)}")
        
        file_name = os.path.basename(file_path).lower()
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # Extract document content if not provided
        if content is None:
            # Extract text manually instead of calling the function tool
            try:
                # Read file based on extension
                if file_ext == ".pdf":
                    try:
                        import fitz
                        doc = fitz.open(file_path)
                        content = ""
                        # Process all pages
                        for page_num in range(len(doc)):
                            page = doc.load_page(page_num)
                            content += page.get_text()
                        doc.close()
                    except ImportError:
                        logger.error("PyMuPDF not installed. Install with: pip install PyMuPDF")
                        return document_type_fallback(file_path, file_name, file_ext)
                    except Exception as e:
                        logger.error(f"Error extracting text from PDF: {str(e)}")
                        return document_type_fallback(file_path, file_name, file_ext)
                
                elif file_ext == ".docx":
                    try:
                        from docx import Document
                        doc = Document(file_path)
                        content = ""
                        # Process all paragraphs
                        for para in doc.paragraphs:
                            content += para.text + "\n"
                    except ImportError:
                        logger.error("python-docx not installed. Install with: pip install python-docx")
                        return document_type_fallback(file_path, file_name, file_ext)
                    except Exception as e:
                        logger.error(f"Error extracting text from DOCX: {str(e)}")
                        return document_type_fallback(file_path, file_name, file_ext)
                
                elif file_ext in [".csv", ".xlsx"]:
                    try:
                        import pandas as pd
                        if file_ext == ".csv":
                            df = pd.read_csv(file_path, on_bad_lines='skip')
                        else:
                            df = pd.read_excel(file_path, engine='openpyxl')
                        content = df.to_string()
                    except ImportError:
                        logger.error("pandas not installed. Install with: pip install pandas openpyxl")
                        return document_type_fallback(file_path, file_name, file_ext)
                    except Exception as e:
                        logger.error(f"Error reading spreadsheet: {str(e)}")
                        return document_type_fallback(file_path, file_name, file_ext)
                
                elif file_ext == ".json":
                    with open(file_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                        content = json.dumps(data, indent=2)
                
                elif file_ext == ".txt":
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                
                else:
                    logger.error(f"Unsupported file extension: {file_ext}")
                    return document_type_fallback(file_path, file_name, file_ext)
                
                # Chunk the content
                chunks = chunk_text(content, config.MAX_TOKENS, config.CHUNK_OVERLAP)
                
            except Exception as e:
                logger.error(f"Error extracting text: {str(e)}")
                return document_type_fallback(file_path, file_name, file_ext)
        else:
            # If content is provided, chunk it
            chunks = chunk_text(content, config.MAX_TOKENS, config.CHUNK_OVERLAP)
        
        logger.info(f"Analyzing document {file_path} with {len(chunks)} chunks in parallel")
        
        # Initialize aggregated metadata
        all_reports = []
        primary_document_type = "unknown"
        primary_confidence = 0.0
        all_detected_fields = []
        primary_metadata = {
            "analyzed_at": datetime.now().isoformat(),
            "report_period": "Unknown",
            "client_name": "Unknown",
            "entity": "Unknown",
            "description": "Unknown document type"
        }
        
        # Process chunks in parallel for speed with timeout
        async def process_chunk(chunk_idx, chunk):
            try:
                prompt = f"""Analyze this document chunk carefully and extract key information for document type detection. 
                            Identify financial metrics, report structure, headings, and any specific terminology:
                            
                            {chunk}
                            
                            Provide detailed analysis of the content:"""
                
                # Increase the timeout from 3 to 10 seconds and add retries
                response = await call_ollama(prompt, config.OLLAMA_MODEL, timeout=10, max_retries=1)
                
                if "Error:" in response:
                    logger.warning(f"Ollama API error or timeout for chunk {chunk_idx+1}: {response}")
                    return None
                
                return response
            except Exception as e:
                logger.error(f"Error processing chunk {chunk_idx+1}: {str(e)}")
                return None
        
        # Process all chunks in parallel
        tasks = [process_chunk(i, chunk) for i, chunk in enumerate(chunks)]
        chunk_results = await asyncio.gather(*tasks)
        
        # Process results
        valid_results = 0
        timeout_count = 0
        for chunk_result in chunk_results:
            if not chunk_result:
                timeout_count += 1
                continue
                
            valid_results += 1
            # Extract document type and confidence from this chunk
            chunk_doc_type = chunk_result.get("primary_document_type", "").lower()
            chunk_confidence = float(chunk_result.get("confidence", 0.0))
            chunk_fields = chunk_result.get("detected_fields", [])
            chunk_metadata = chunk_result.get("metadata", {})
            detected_reports = chunk_result.get("detected_reports", [])
            
            # Extend all_reports with any new reports found in this chunk
            for report in detected_reports:
                if report not in all_reports:
                    all_reports.append(report)
            
            # Update the primary document type if this chunk has higher confidence
            if chunk_confidence > primary_confidence:
                primary_document_type = chunk_doc_type
                primary_confidence = chunk_confidence
                primary_metadata = chunk_metadata
            
            # Add unique detected fields
            for field in chunk_fields:
                if field not in all_detected_fields:
                    all_detected_fields.append(field)
            
            # If client_name is still unknown but found in this chunk, update it
            if primary_metadata.get("client_name") == "Unknown" and chunk_metadata.get("client_name"):
                primary_metadata["client_name"] = chunk_metadata.get("client_name")
            
            # If report_period is still unknown but found in this chunk, update it
            if primary_metadata.get("report_period") == "Unknown" and chunk_metadata.get("report_period"):
                primary_metadata["report_period"] = chunk_metadata.get("report_period")
        
        # If most chunks timed out, use fast analysis instead of fallback
        if timeout_count > len(chunks) * 0.5:
            logger.warning(f"Majority of chunks timed out ({timeout_count}/{len(chunks)}), using fast analysis")
            fast_result = fast_analyze_document_type(file_path)
            
            # If we got some valid results, combine them with fast analysis
            if valid_results > 0:
                # Keep the document type from our valid results if we have high confidence
                if primary_confidence > 0.7:
                    fast_result.document_type = primary_document_type
                    fast_result.confidence = primary_confidence
                
                # Add any detected fields from valid results
                for field in all_detected_fields:
                    if field not in fast_result.detected_fields:
                        fast_result.detected_fields.append(field)
                
                # Combine metadata
                for key, value in primary_metadata.items():
                    if value != "Unknown" and (key not in fast_result.metadata or fast_result.metadata[key] == "Unknown"):
                        fast_result.metadata[key] = value
                
                fast_result.metadata["analysis_method"] = "hybrid"
            
            # Cache the result
            if file_hash:
                try:
                    cache_dir = os.path.join(config.DATA_DIR, 'cache')
                    os.makedirs(cache_dir, exist_ok=True)
                    cache_file = os.path.join(cache_dir, f"{file_hash}.json")
                    
                    # Cache the analysis result
                    cache_data = {
                        "document_type": fast_result.document_type,
                        "confidence": fast_result.confidence,
                        "detected_fields": fast_result.detected_fields,
                        "metadata": fast_result.metadata
                    }
                    
                    with open(cache_file, 'w') as f:
                        json.dump(cache_data, f)
                    
                    logger.info(f"Cached analysis result for {file_path}")
                except Exception as e:
                    logger.warning(f"Couldn't cache analysis for {file_path}: {str(e)}")
            
            return fast_result
        
        # If no valid results were processed, use fallback
        if valid_results == 0:
            logger.warning(f"No valid analysis results for {file_path}, using fallback")
            return document_type_fallback(file_path, file_name, file_ext)
        
        # Ensure document_type is one of the recognized types
        recognized_types = [
            "annual_report", "audit_report", "balance_sheet", 
            "income_statement", "cash_flow_statement", "tax_document", 
            "bank_statement", "financial_report"
        ]
        
        if primary_document_type not in recognized_types:
            # Try to map to closest match
            if "balance" in primary_document_type:
                primary_document_type = "balance_sheet"
            elif "income" in primary_document_type:
                primary_document_type = "income_statement"
            elif "cash" in primary_document_type:
                primary_document_type = "cash_flow_statement"
            elif "tax" in primary_document_type:
                primary_document_type = "tax_document"
            elif "audit" in primary_document_type:
                primary_document_type = "audit_report"
            elif "annual" in primary_document_type or "report" in primary_document_type:
                primary_document_type = "annual_report"
            elif "bank" in primary_document_type:
                primary_document_type = "bank_statement"
            else:
                primary_document_type = "financial_report"
        
        # If we didn't find any reports or if analysis failed, use filename-based fallback
        if primary_document_type == "unknown" or primary_confidence < 0.3:
            logger.warning(f"Analysis of {file_path} produced low confidence results, using fallback")
            fallback = document_type_fallback(file_path, file_name, file_ext)
            primary_document_type = fallback.document_type
            primary_confidence = fallback.confidence
            
            # Only use fallback metadata if we don't have any better metadata
            if primary_metadata.get("client_name") == "Unknown":
                primary_metadata = fallback.metadata
        
        # Add report types to metadata if detected
        if all_reports:
            primary_metadata["detected_reports"] = all_reports
            
            # Add report summary to metadata
            report_types = [report.get("report_type") for report in all_reports if report.get("report_type")]
            if report_types:
                primary_metadata["report_types_summary"] = report_types
        
        # Create the analysis output
        result = DocumentTypeAnalysisOutput(
            file_path=file_path,
            document_type=primary_document_type,
            confidence=primary_confidence,
            detected_fields=all_detected_fields,
            metadata=primary_metadata
        )
        
        # Cache the result for future use if we have a file hash
        if file_hash:
            try:
                cache_dir = os.path.join(config.DATA_DIR, 'cache')
                os.makedirs(cache_dir, exist_ok=True)
                cache_file = os.path.join(cache_dir, f"{file_hash}.json")
                
                # Cache the analysis result
                cache_data = {
                    "document_type": primary_document_type,
                    "confidence": primary_confidence,
                    "detected_fields": all_detected_fields,
                    "metadata": primary_metadata
                }
                
                with open(cache_file, 'w') as f:
                    json.dump(cache_data, f)
                
                logger.info(f"Cached analysis result for {file_path}")
            except Exception as e:
                logger.warning(f"Couldn't cache analysis for {file_path}: {str(e)}")
        
        return result
            
    except Exception as e:
        logger.error(f"Error in analyze_document_type: {str(e)}")
        # Return basic fallback with error
        return DocumentTypeAnalysisOutput(
            file_path=file_path,
            document_type="unknown",
            confidence=0.1,
            detected_fields=[],
            metadata={
                "analyzed_at": datetime.now().isoformat(),
                "error": str(e)
            }
        )

def document_type_fallback(file_path: str, file_name: str, file_ext: str) -> DocumentTypeAnalysisOutput:
    """Fallback method to determine document type from file name and extension"""
    try:
        # Use the fast analysis method instead of simple fallback
        return fast_analyze_document_type(file_path)
    except Exception as e:
        # If fast_analyze_document_type fails, provide a very basic fallback
        logger.error(f"Error in fast_analyze_document_type: {str(e)}")
        file_id = abs(hash(file_path)) % 10000
        
        # Check if file exists - if not, provide appropriate metadata
        if not os.path.exists(file_path):
            return DocumentTypeAnalysisOutput(
                file_path=file_path,
                document_type="unknown",
                confidence=0.1,
                detected_fields=["error"],
                metadata={
                    "analyzed_at": datetime.now().isoformat(),
                    "report_period": "Unknown",
                    "client_name": "Unknown", 
                    "entity": "Unknown",
                    "account_name": "Unknown",
                    "wallet_id": f"WLT_ERR_{file_id:04d}",
                    "description": f"File not found: {file_path}",
                    "analysis_method": "emergency_fallback",
                    "error": f"File not found: {str(e)}"
                }
            )
        
        # Basic document type inference from filename for robust fallback
        document_type = "financial_report"  # Default
        detected_fields = ["document"]
        description = f"Financial document (fallback analysis)"
        
        # Simple filename-based inference
        file_name_lower = file_name.lower()
        if "balance" in file_name_lower or "sheet" in file_name_lower:
            document_type = "balance_sheet"
            detected_fields = ["financial_data"]
            description = "Balance sheet (fallback analysis)"
        elif "income" in file_name_lower or "profit" in file_name_lower:
            document_type = "income_statement"
            detected_fields = ["financial_data"]
            description = "Income statement (fallback analysis)"
        elif "cash" in file_name_lower or "flow" in file_name_lower:
            document_type = "cash_flow_statement" 
            detected_fields = ["financial_data"]
            description = "Cash flow statement (fallback analysis)"
        elif "tax" in file_name_lower:
            document_type = "tax_document"
            detected_fields = ["financial_data"]
            description = "Tax document (fallback analysis)"
        elif "bank" in file_name_lower or "statement" in file_name_lower:
            document_type = "bank_statement"
            detected_fields = ["financial_data"]
            description = "Bank statement (fallback analysis)"
        elif "annual" in file_name_lower or "report" in file_name_lower:
            document_type = "annual_report"
            detected_fields = ["financial_data"]
            description = "Annual report (fallback analysis)"
        
        return DocumentTypeAnalysisOutput(
            file_path=file_path,
            document_type=document_type,
            confidence=0.5,
            detected_fields=detected_fields,
            metadata={
                "analyzed_at": datetime.now().isoformat(),
                "report_period": "Unknown",
                "client_name": "Unknown",
                "entity": "Unknown",
                "account_name": "Unknown",
                "wallet_id": f"WLT_{file_id:04d}",
                "description": description,
                "analysis_method": "basic_fallback",
                "error": str(e)
            }
        )

class DocumentAnalysisOutput(StrictBaseModel):
    status: str
    message: Optional[str] = None
    file_info: Optional[Dict[str, Any]] = None
    document_type: str = "unknown"
    extraction_time: str = ""

@function_tool
async def analyze_document(file_path: str) -> DocumentAnalysisOutput:
    """
    Analyze a document to extract its content and structure
    
    Args:
        file_path: Path to the document
        
    Returns:
        Document analysis results
    """
    logger.info(f"Analyzing document: {file_path}")
    
    try:
        file_info = get_file_info(file_path)
        
        return DocumentAnalysisOutput(
            status="success",
            message=f"Document {file_path} analyzed",
            file_info=file_info,
            document_type="unknown",
            extraction_time=datetime.now().isoformat()
        )
    except Exception as e:
        logger.error(f"Error analyzing document: {str(e)}")
        return DocumentAnalysisOutput(
            status="error",
            message=f"Error analyzing document: {str(e)}"
        )

async def extract_annual_report_data(file_path: str, content: Optional[str] = None) -> Dict[str, Any]:
    """Extract data from an annual report"""
    if not content:
        try:
            # Use the document text extraction logic
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == ".pdf":
                import fitz
                doc = fitz.open(file_path)
                content = ""
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    content += page.get_text()
                doc.close()
            elif file_ext == ".docx":
                from docx import Document
                doc = Document(file_path)
                content = ""
                for para in doc.paragraphs:
                    content += para.text + "\n"
            elif file_ext in [".csv", ".xlsx"]:
                import pandas as pd
                if file_ext == ".csv":
                    df = pd.read_csv(file_path, on_bad_lines='skip')
                else:
                    df = pd.read_excel(file_path, engine='openpyxl')
                content = df.to_string()
            elif file_ext == ".json":
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    content = json.dumps(data, indent=2)
            elif file_ext == ".txt":
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
        except Exception as e:
            logger.error(f"Error extracting content for annual report: {str(e)}")
            content = ""
    
    # Quick extraction without Ollama to avoid timeouts
    result = {
        "report_type": "annual_report",
        "extracted_at": datetime.now().isoformat(),
        "data_points": {}
    }
    
    # Extract basic data by simple pattern matching
    if "revenue" in content.lower():
        result["data_points"]["has_revenue"] = True
    if "profit" in content.lower():
        result["data_points"]["has_profit"] = True
    if "loss" in content.lower():
        result["data_points"]["has_loss"] = True
    if "assets" in content.lower():
        result["data_points"]["has_assets"] = True
    if "liabilities" in content.lower():
        result["data_points"]["has_liabilities"] = True
    
    return result

async def extract_audit_report_data(file_path: str, content: Optional[str] = None) -> Dict[str, Any]:
    """Extract data from an audit report"""
    if not content:
        try:
            # Use the document text extraction logic
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == ".pdf":
                import fitz
                doc = fitz.open(file_path)
                content = ""
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    content += page.get_text()
                doc.close()
            elif file_ext == ".docx":
                from docx import Document
                doc = Document(file_path)
                content = ""
                for para in doc.paragraphs:
                    content += para.text + "\n"
            elif file_ext in [".csv", ".xlsx"]:
                import pandas as pd
                if file_ext == ".csv":
                    df = pd.read_csv(file_path, on_bad_lines='skip')
                else:
                    df = pd.read_excel(file_path, engine='openpyxl')
                content = df.to_string()
            elif file_ext == ".json":
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    content = json.dumps(data, indent=2)
            elif file_ext == ".txt":
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
        except Exception as e:
            logger.error(f"Error extracting content for audit report: {str(e)}")
            content = ""
    
    # Quick extraction without Ollama to avoid timeouts
    result = {
        "report_type": "audit_report",
        "extracted_at": datetime.now().isoformat(),
        "data_points": {}
    }
    
    # Extract basic data by simple pattern matching
    if "findings" in content.lower():
        result["data_points"]["has_findings"] = True
    if "compliance" in content.lower():
        result["data_points"]["has_compliance"] = True
    if "recommendation" in content.lower():
        result["data_points"]["has_recommendations"] = True
    if "control" in content.lower():
        result["data_points"]["has_controls"] = True
    
    return result

async def extract_balance_sheet_data(file_path: str, content: Optional[str] = None) -> Dict[str, Any]:
    """Extract data from a balance sheet"""
    if not content:
        try:
            # Use the document text extraction logic
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == ".pdf":
                import fitz
                doc = fitz.open(file_path)
                content = ""
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    content += page.get_text()
                doc.close()
            elif file_ext == ".docx":
                from docx import Document
                doc = Document(file_path)
                content = ""
                for para in doc.paragraphs:
                    content += para.text + "\n"
            elif file_ext in [".csv", ".xlsx"]:
                import pandas as pd
                if file_ext == ".csv":
                    df = pd.read_csv(file_path, on_bad_lines='skip')
                else:
                    df = pd.read_excel(file_path, engine='openpyxl')
                content = df.to_string()
            elif file_ext == ".json":
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    content = json.dumps(data, indent=2)
            elif file_ext == ".txt":
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
        except Exception as e:
            logger.error(f"Error extracting content for balance sheet: {str(e)}")
            content = ""
    
    # Quick extraction without Ollama to avoid timeouts
    result = {
        "report_type": "balance_sheet",
        "extracted_at": datetime.now().isoformat(),
        "data_points": {}
    }
    
    # Extract basic data by simple pattern matching
    if "assets" in content.lower():
        result["data_points"]["has_assets"] = True
    if "liabilities" in content.lower():
        result["data_points"]["has_liabilities"] = True
    if "equity" in content.lower():
        result["data_points"]["has_equity"] = True
    if "cash" in content.lower():
        result["data_points"]["has_cash"] = True
    
    return result

async def extract_income_statement_data(file_path: str, content: Optional[str] = None) -> Dict[str, Any]:
    """Extract data from an income statement"""
    if not content:
        try:
            # Use the document text extraction logic
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == ".pdf":
                import fitz
                doc = fitz.open(file_path)
                content = ""
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    content += page.get_text()
                doc.close()
            elif file_ext == ".docx":
                from docx import Document
                doc = Document(file_path)
                content = ""
                for para in doc.paragraphs:
                    content += para.text + "\n"
            elif file_ext in [".csv", ".xlsx"]:
                import pandas as pd
                if file_ext == ".csv":
                    df = pd.read_csv(file_path, on_bad_lines='skip')
                else:
                    df = pd.read_excel(file_path, engine='openpyxl')
                content = df.to_string()
            elif file_ext == ".json":
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    content = json.dumps(data, indent=2)
            elif file_ext == ".txt":
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
        except Exception as e:
            logger.error(f"Error extracting content for income statement: {str(e)}")
            content = ""
    
    # Quick extraction without Ollama to avoid timeouts
    result = {
        "report_type": "income_statement",
        "extracted_at": datetime.now().isoformat(),
        "data_points": {}
    }
    
    # Extract basic data by simple pattern matching
    if "revenue" in content.lower():
        result["data_points"]["has_revenue"] = True
    if "expense" in content.lower():
        result["data_points"]["has_expenses"] = True
    if "income" in content.lower():
        result["data_points"]["has_income"] = True
    if "profit" in content.lower():
        result["data_points"]["has_profit"] = True
    if "tax" in content.lower():
        result["data_points"]["has_tax"] = True
    
    return result

async def extract_cash_flow_statement_data(file_path: str, content: Optional[str] = None) -> Dict[str, Any]:
    """Extract data from a cash flow statement"""
    if not content:
        try:
            # Use the document text extraction logic
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == ".pdf":
                import fitz
                doc = fitz.open(file_path)
                content = ""
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    content += page.get_text()
                doc.close()
            elif file_ext == ".docx":
                from docx import Document
                doc = Document(file_path)
                content = ""
                for para in doc.paragraphs:
                    content += para.text + "\n"
            elif file_ext in [".csv", ".xlsx"]:
                import pandas as pd
                if file_ext == ".csv":
                    df = pd.read_csv(file_path, on_bad_lines='skip')
                else:
                    df = pd.read_excel(file_path, engine='openpyxl')
                content = df.to_string()
            elif file_ext == ".json":
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    content = json.dumps(data, indent=2)
            elif file_ext == ".txt":
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
        except Exception as e:
            logger.error(f"Error extracting content for cash flow statement: {str(e)}")
            content = ""
    
    # Quick extraction without Ollama to avoid timeouts
    result = {
        "report_type": "cash_flow_statement",
        "extracted_at": datetime.now().isoformat(),
        "data_points": {}
    }
    
    # Extract basic data by simple pattern matching
    if "operating" in content.lower():
        result["data_points"]["has_operating"] = True
    if "investing" in content.lower():
        result["data_points"]["has_investing"] = True
    if "financing" in content.lower():
        result["data_points"]["has_financing"] = True
    if "cash" in content.lower():
        result["data_points"]["has_cash"] = True
    
    return result

async def extract_tax_document_data(file_path: str, content: Optional[str] = None) -> Dict[str, Any]:
    """Extract data from a tax document"""
    if not content:
        try:
            # Use the document text extraction logic
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == ".pdf":
                import fitz
                doc = fitz.open(file_path)
                content = ""
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    content += page.get_text()
                doc.close()
            elif file_ext == ".docx":
                from docx import Document
                doc = Document(file_path)
                content = ""
                for para in doc.paragraphs:
                    content += para.text + "\n"
            elif file_ext in [".csv", ".xlsx"]:
                import pandas as pd
                if file_ext == ".csv":
                    df = pd.read_csv(file_path, on_bad_lines='skip')
                else:
                    df = pd.read_excel(file_path, engine='openpyxl')
                content = df.to_string()
            elif file_ext == ".json":
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    content = json.dumps(data, indent=2)
            elif file_ext == ".txt":
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
        except Exception as e:
            logger.error(f"Error extracting content for tax document: {str(e)}")
            content = ""
    
    # Quick extraction without Ollama to avoid timeouts
    result = {
        "report_type": "tax_document",
        "extracted_at": datetime.now().isoformat(),
        "data_points": {}
    }
    
    # Extract basic data by simple pattern matching
    if "tax" in content.lower():
        result["data_points"]["has_tax"] = True
    if "deduction" in content.lower():
        result["data_points"]["has_deductions"] = True
    if "credit" in content.lower():
        result["data_points"]["has_credits"] = True
    if "payment" in content.lower():
        result["data_points"]["has_payments"] = True
    
    return result

async def extract_bank_statement_data(file_path: str, content: Optional[str] = None) -> Dict[str, Any]:
    """Extract data from a bank statement"""
    if not content:
        try:
            # Use the document text extraction logic
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == ".pdf":
                import fitz
                doc = fitz.open(file_path)
                content = ""
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    content += page.get_text()
                doc.close()
            elif file_ext == ".docx":
                from docx import Document
                doc = Document(file_path)
                content = ""
                for para in doc.paragraphs:
                    content += para.text + "\n"
            elif file_ext in [".csv", ".xlsx"]:
                import pandas as pd
                if file_ext == ".csv":
                    df = pd.read_csv(file_path, error_bad_lines=False, warn_bad_lines=True)
                else:
                    df = pd.read_excel(file_path, engine='openpyxl')
                content = df.to_string()
            elif file_ext == ".json":
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    content = json.dumps(data, indent=2)
            elif file_ext == ".txt":
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
        except Exception as e:
            logger.error(f"Error extracting content for bank statement: {str(e)}")
            content = ""
    
    # Quick extraction without Ollama to avoid timeouts
    result = {
        "report_type": "bank_statement",
        "extracted_at": datetime.now().isoformat(),
        "data_points": {}
    }
    
    # Extract basic data by simple pattern matching
    if "transaction" in content.lower():
        result["data_points"]["has_transactions"] = True
    if "deposit" in content.lower():
        result["data_points"]["has_deposits"] = True
    if "withdrawal" in content.lower():
        result["data_points"]["has_withdrawals"] = True
    if "balance" in content.lower():
        result["data_points"]["has_balance"] = True
    
    return result

def fast_analyze_document_type(file_path: str) -> DocumentTypeAnalysisOutput:
    """
    Quickly analyze a document's type using heuristics and filename patterns
    without relying on Ollama. This provides an immediate result when speed is critical.
    
    Args:
        file_path: Path to the document
        
    Returns:
        Document type information
    """
    # Convert Path object to string if needed
    if not isinstance(file_path, str):
        file_path = str(file_path)
    
    # Check if file exists
    if not os.path.exists(file_path):
        logger.error(f"File not found in fast_analyze_document_type: {file_path}")
        # Create a fallback document type analysis for non-existent files
        return DocumentTypeAnalysisOutput(
            file_path=file_path,
            document_type="unknown",
            confidence=0.1,
            detected_fields=["error"],
            metadata={
                "analyzed_at": datetime.now().isoformat(),
                "report_period": "Unknown",
                "client_name": "Unknown",
                "entity": "Unknown",
                "account_name": "Unknown",
                "wallet_id": f"WLT_ERR",
                "description": f"File not found: {file_path}",
                "analysis_method": "error",
                "error": "File not found"
            }
        )
    
    file_name = os.path.basename(file_path).lower()
    file_ext = os.path.splitext(file_path)[1].lower()
    
    # Generate a simple ID for the document
    file_id = abs(hash(file_path)) % 10000
    
    # Default metadata
    metadata = {
        "analyzed_at": datetime.now().isoformat(),
        "report_period": "2023-2024",
        "client_name": "ACME Corporation",
        "entity": "Finance Department",
        "account_name": "General",
        "wallet_id": f"WLT_{file_id:04d}",
        "description": "Financial document",
        "analysis_method": "fast"
    }
    
    # Determine document type based on filename keywords first
    document_type = "financial_report"  # Default to a generic type instead of "unknown"
    confidence = 0.6
    detected_fields = ["financial_data"]
    
    # Try to check file content for more accurate type determination if it's a text-based file
    try:
        # For PDFs, try to read first page to confirm it's valid
        if file_ext == ".pdf":
            try:
                import fitz
                doc = fitz.open(file_path)
                if len(doc) > 0:
                    # Read first page content to help determine type
                    page = doc.load_page(0)
                    content = page.get_text(sort=True)
                    doc.close()
                    
                    # Use content to refine document type
                    if "balance sheet" in content.lower():
                        document_type = "balance_sheet"
                        confidence = 0.95
                        detected_fields = ["assets", "liabilities", "equity"]
                        metadata["description"] = "Statement of financial position"
                    elif "income statement" in content.lower() or "profit and loss" in content.lower():
                        document_type = "income_statement"
                        confidence = 0.95
                        detected_fields = ["revenue", "expenses", "profit"]
                        metadata["description"] = "Statement of profit and loss"
                    elif "cash flow" in content.lower():
                        document_type = "cash_flow_statement"
                        confidence = 0.95
                        detected_fields = ["operating_activities", "investing_activities", "financing_activities"]
                        metadata["description"] = "Statement of cash flows"
                    elif "tax" in content.lower() and ("form" in content.lower() or "return" in content.lower()):
                        document_type = "tax_document"
                        confidence = 0.95
                        detected_fields = ["tax_liability", "deductions", "tax_credits"]
                        metadata["description"] = "Tax documentation"
                    elif "bank statement" in content.lower() or "account statement" in content.lower():
                        document_type = "bank_statement"
                        confidence = 0.95
                        detected_fields = ["transactions", "balance", "account_details"]
                        metadata["description"] = "Bank account statement"
                        metadata["account_name"] = "Primary Account"
            except Exception as e:
                logger.warning(f"Could not analyze PDF content: {str(e)}")
                # Continue with filename-based analysis
        
        # For CSV files, check headers to determine type
        elif file_ext == ".csv":
            try:
                import pandas as pd
                try:
                    df = pd.read_csv(file_path, nrows=1, on_bad_lines='skip')
                    headers = [col.lower() for col in df.columns]
                    
                    # Use headers to determine document type
                    if any(term in headers for term in ["asset", "liability", "equity"]):
                        document_type = "balance_sheet"
                        confidence = 0.95
                        detected_fields = ["assets", "liabilities", "equity"]
                        metadata["description"] = "Statement of financial position"
                    elif any(term in headers for term in ["revenue", "income", "expense", "profit", "loss"]):
                        document_type = "income_statement"
                        confidence = 0.95
                        detected_fields = ["revenue", "expenses", "profit"]
                        metadata["description"] = "Statement of profit and loss"
                    elif any(term in headers for term in ["cash", "flow", "operating", "investing", "financing"]):
                        document_type = "cash_flow_statement"
                        confidence = 0.95
                        detected_fields = ["operating_activities", "investing_activities", "financing_activities"]
                        metadata["description"] = "Statement of cash flows"
                    elif any(term in headers for term in ["tax", "deduction", "credit"]):
                        document_type = "tax_document"
                        confidence = 0.95
                        detected_fields = ["tax_liability", "deductions", "tax_credits"]
                        metadata["description"] = "Tax documentation"
                    elif any(term in headers for term in ["transaction", "date", "amount", "balance", "account"]):
                        document_type = "bank_statement"
                        confidence = 0.95
                        detected_fields = ["transactions", "balance", "account_details"]
                        metadata["description"] = "Bank account statement"
                        # Try to extract account name from data
                        try:
                            if "account" in headers:
                                account_idx = headers.index("account")
                                acct_value = df.iloc[0, account_idx]
                                if isinstance(acct_value, str) and acct_value.strip():
                                    metadata["account_name"] = acct_value.strip()
                        except:
                            pass
                except pd.errors.EmptyDataError:
                    logger.warning(f"CSV file is empty: {file_path}")
                except Exception as e:
                    logger.warning(f"Error reading CSV file: {str(e)}")
            except Exception as e:
                logger.warning(f"Could not analyze CSV headers: {str(e)}")
                # Continue with filename-based analysis
    except Exception as e:
        logger.warning(f"Content-based analysis failed: {str(e)}")
        # Continue with filename-based analysis
    
    # Check keywords in filename if we still have default confidence
    if confidence <= 0.6:
        if "annual" in file_name or "report" in file_name:
            document_type = "annual_report"
            confidence = 0.8
            detected_fields = ["company_performance", "financial_summary", "outlook"]
            metadata["report_period"] = "FY 2023"
            metadata["description"] = "Annual company performance report"
        elif "audit" in file_name:
            document_type = "audit_report"
            confidence = 0.8
            detected_fields = ["findings", "recommendations", "compliance"]
            metadata["report_period"] = "Q4 2023"
            metadata["description"] = "Independent auditor's assessment"
        elif "balance" in file_name or "sheet" in file_name:
            document_type = "balance_sheet"
            confidence = 0.9
            detected_fields = ["assets", "liabilities", "equity"]
            metadata["report_period"] = "Q4 2023"
            metadata["description"] = "Statement of financial position"
        elif "income" in file_name or "statement" in file_name or "profit" in file_name:
            document_type = "income_statement"
            confidence = 0.9
            detected_fields = ["revenue", "expenses", "profit"]
            metadata["report_period"] = "Q4 2023"
            metadata["description"] = "Statement of profit and loss"
        elif "cash" in file_name or "flow" in file_name:
            document_type = "cash_flow_statement"
            confidence = 0.9
            detected_fields = ["operating_activities", "investing_activities", "financing_activities"]
            metadata["report_period"] = "Q4 2023"
            metadata["description"] = "Statement of cash flows"
        elif "tax" in file_name:
            document_type = "tax_document"
            confidence = 0.8
            detected_fields = ["tax_liability", "deductions", "tax_credits"]
            metadata["report_period"] = "Tax Year 2023"
            metadata["description"] = "Tax documentation"
        elif "bank" in file_name:
            document_type = "bank_statement"
            confidence = 0.9
            detected_fields = ["transactions", "balance", "account_details"]
            metadata["report_period"] = "Monthly Statement"
            metadata["description"] = "Bank account statement"
            metadata["account_name"] = "Primary Account"
    
    # Return the document type analysis output
    return DocumentTypeAnalysisOutput(
        file_path=file_path,
        document_type=document_type,
        confidence=confidence,
        detected_fields=detected_fields,
        metadata=metadata
    )

async def extract_data_using_ollama_agent(combined_text, doc_type, file_name):
    """
    Extract financial data from text using Ollama as an agent-based approach
    
    Args:
        combined_text: Text to extract data from
        doc_type: Type of document (balance_sheet, income_statement, etc.)
        file_name: Name of the file being processed
        
    Returns:
        Structured data for CSV generation
    """
    logger.info(f"Extracting data using Ollama agent approach for {file_name}")
    
    # Create a more specific prompt based on document type
    if doc_type == "balance_sheet":
        prompt_template = """
You are a financial data extraction specialist. Extract all balance sheet data from this document into a structured format.

DOCUMENT: {file_name}
CONTENT:
{text}

Extract every financial line item with the following structure:
{
  "columns": ["Account", "Amount", "Category", "Type", "Notes"],
  "data": [
    {"Account": "Cash and Cash Equivalents", "Amount": "250000", "Category": "Assets", "Type": "Current Asset", "Notes": "Cash in bank accounts"},
    {"Account": "Accounts Receivable", "Amount": "120000", "Category": "Assets", "Type": "Current Asset", "Notes": "Outstanding customer payments"},
    {"Account": "Property and Equipment", "Amount": "500000", "Category": "Assets", "Type": "Fixed Asset", "Notes": "Office building and equipment"},
    {"Account": "Accounts Payable", "Amount": "80000", "Category": "Liabilities", "Type": "Current Liability", "Notes": "Vendor invoices due"},
    {"Account": "Long-term Debt", "Amount": "300000", "Category": "Liabilities", "Type": "Long-term Liability", "Notes": "Bank loan"},
    {"Account": "Common Stock", "Amount": "200000", "Category": "Equity", "Type": "Equity", "Notes": "Shareholder investments"},
    {"Account": "Retained Earnings", "Amount": "290000", "Category": "Equity", "Type": "Equity", "Notes": "Accumulated profits"}
  ]
}

IMPORTANT INSTRUCTIONS:
1. Extract EVERY line item from the balance sheet with its exact amount
2. Categorize each item as Assets, Liabilities, or Equity
3. Identify the type of each item (Current, Fixed, Long-term, etc.)
4. Include all assets, liabilities, and equity accounts
5. Format numbers as plain numbers without currency symbols or commas
6. Ensure the data is structured in a way that matches the original document hierarchy

EXTRACT ONLY THE JSON:"""
    
    elif doc_type == "income_statement":
        prompt_template = """
You are a financial data extraction specialist. Extract all income statement data from this document into a structured format.

DOCUMENT: {file_name}
CONTENT:
{text}

Extract every financial line item with the following structure:
{
  "columns": ["Item", "Amount", "Category", "Period", "Notes"],
  "data": [
    {"Item": "Sales Revenue", "Amount": "850000", "Category": "Revenue", "Period": "Q1 2023", "Notes": "Product sales"},
    {"Item": "Service Revenue", "Amount": "250000", "Category": "Revenue", "Period": "Q1 2023", "Notes": "Consulting services"},
    {"Item": "Cost of Goods Sold", "Amount": "400000", "Category": "Expense", "Period": "Q1 2023", "Notes": "Production costs"},
    {"Item": "Salaries and Wages", "Amount": "300000", "Category": "Expense", "Period": "Q1 2023", "Notes": "Employee compensation"},
    {"Item": "Rent Expense", "Amount": "50000", "Category": "Expense", "Period": "Q1 2023", "Notes": "Office space"},
    {"Item": "Depreciation", "Amount": "25000", "Category": "Expense", "Period": "Q1 2023", "Notes": "Equipment depreciation"},
    {"Item": "Net Income", "Amount": "325000", "Category": "Profit", "Period": "Q1 2023", "Notes": "Total profit for period"}
  ]
}

IMPORTANT INSTRUCTIONS:
1. Extract EVERY line item from the income statement with its exact amount
2. Categorize each item as Revenue, Expense, or Profit
3. Identify the time period for each item if available
4. Include all revenue sources, expense categories, and profit figures
5. Format numbers as plain numbers without currency symbols or commas
6. Ensure the data is structured in a way that matches the original document hierarchy

EXTRACT ONLY THE JSON:"""
    
    elif doc_type == "cash_flow":
        prompt_template = """
You are a financial data extraction specialist. Extract all cash flow statement data from this document into a structured format.

DOCUMENT: {file_name}
CONTENT:
{text}

Extract every financial line item with the following structure:
{
  "columns": ["Activity", "Item", "Amount", "Category", "Notes"],
  "data": [
    {"Activity": "Operating", "Item": "Net Income", "Amount": "325000", "Category": "Inflow", "Notes": "From income statement"},
    {"Activity": "Operating", "Item": "Depreciation", "Amount": "25000", "Category": "Adjustment", "Notes": "Non-cash expense"},
    {"Activity": "Operating", "Item": "Increase in Accounts Receivable", "Amount": "-45000", "Category": "Outflow", "Notes": "More outstanding payments"},
    {"Activity": "Investing", "Item": "Purchase of Equipment", "Amount": "-120000", "Category": "Outflow", "Notes": "New manufacturing equipment"},
    {"Activity": "Financing", "Item": "Loan Proceeds", "Amount": "200000", "Category": "Inflow", "Notes": "New bank loan"},
    {"Activity": "Financing", "Item": "Dividend Payment", "Amount": "-50000", "Category": "Outflow", "Notes": "Quarterly dividend"}
  ]
}

IMPORTANT INSTRUCTIONS:
1. Extract EVERY line item from the cash flow statement with its exact amount
2. Categorize each activity as Operating, Investing, or Financing
3. Identify whether each item is an Inflow, Outflow, or Adjustment
4. Include all cash flow activities and their specific items
5. Format numbers as plain numbers without currency symbols (use negative numbers for outflows)
6. Ensure the data is structured in a way that matches the original document hierarchy

EXTRACT ONLY THE JSON:"""
    
    elif doc_type == "bank_statement":
        prompt_template = """
You are a financial data extraction specialist. Extract all bank statement transaction data from this document into a structured format.

DOCUMENT: {file_name}
CONTENT:
{text}

Extract every transaction with the following structure:
{
  "columns": ["Date", "Description", "Amount", "Type", "Balance"],
  "data": [
    {"Date": "2023-01-01", "Description": "Beginning Balance", "Amount": "", "Type": "", "Balance": "5000.00"},
    {"Date": "2023-01-05", "Description": "Payroll Deposit", "Amount": "2500.00", "Type": "Credit", "Balance": "7500.00"},
    {"Date": "2023-01-12", "Description": "Rent Payment", "Amount": "-1500.00", "Type": "Debit", "Balance": "6000.00"},
    {"Date": "2023-01-15", "Description": "Utility Bill", "Amount": "-125.00", "Type": "Debit", "Balance": "5875.00"},
    {"Date": "2023-01-20", "Description": "Client Payment", "Amount": "3000.00", "Type": "Credit", "Balance": "8875.00"},
    {"Date": "2023-01-31", "Description": "Ending Balance", "Amount": "", "Type": "", "Balance": "8875.00"}
  ]
}

IMPORTANT INSTRUCTIONS:
1. Extract EVERY transaction from the bank statement with its exact date, description and amount
2. Categorize each transaction as Credit (deposit) or Debit (withdrawal)
3. Include the running balance for each transaction if available
4. Format dates consistently (YYYY-MM-DD format if possible)
5. Format amounts as plain numbers (use negative numbers for debits/withdrawals)
6. List transactions in chronological order

EXTRACT ONLY THE JSON:"""
    
    elif doc_type == "annual_report":
        prompt_template = """
You are a financial data extraction specialist. Extract all key financial data from this annual report into a structured format.

DOCUMENT: {file_name}
CONTENT:
{text}

Extract the key financial metrics with the following structure:
{
  "columns": ["Section", "Item", "Amount", "Year", "Notes"],
  "data": [
    {"Section": "Financial Highlights", "Item": "Total Revenue", "Amount": "5200000", "Year": "2023", "Notes": "15% increase YoY"},
    {"Section": "Financial Highlights", "Item": "Net Income", "Amount": "1200000", "Year": "2023", "Notes": "12% increase YoY"},
    {"Section": "Financial Highlights", "Item": "Earnings Per Share", "Amount": "2.15", "Year": "2023", "Notes": "Basic EPS"},
    {"Section": "Balance Sheet", "Item": "Total Assets", "Amount": "7500000", "Year": "2023", "Notes": "As of December 31"},
    {"Section": "Balance Sheet", "Item": "Total Liabilities", "Amount": "3800000", "Year": "2023", "Notes": "As of December 31"},
    {"Section": "Income Statement", "Item": "Operating Expenses", "Amount": "2800000", "Year": "2023", "Notes": "Includes R&D and marketing"},
    {"Section": "Cash Flow", "Item": "Operating Cash Flow", "Amount": "1600000", "Year": "2023", "Notes": "Strong operational performance"}
  ]
}

IMPORTANT INSTRUCTIONS:
1. Extract ALL key financial metrics mentioned in the annual report
2. Organize data by sections (Financial Highlights, Balance Sheet, Income Statement, Cash Flow, etc.)
3. Include the fiscal year for each metric
4. Add relevant notes or context for each metric
5. Format numbers as plain numbers without currency symbols or commas
6. Include all significant metrics that would be useful for financial analysis

EXTRACT ONLY THE JSON:"""
    
    else:
        # Generic financial data format for any other document type
        prompt_template = """
You are a financial data extraction specialist. Extract all financial data from this document into a structured format.

DOCUMENT: {file_name}
CONTENT:
{text}

Extract every financial data point with the following structure:
{
  "columns": ["Category", "Item", "Amount", "Period", "Notes"],
  "data": [
    {"Category": "Revenue", "Item": "Total Revenue", "Amount": "1100000", "Period": "2023", "Notes": "Combined product and service revenue"},
    {"Category": "Expense", "Item": "Operating Expenses", "Amount": "750000", "Period": "2023", "Notes": "All operating costs"},
    {"Category": "Profit", "Item": "Net Income", "Amount": "350000", "Period": "2023", "Notes": "Profit after all expenses"},
    {"Category": "Assets", "Item": "Total Assets", "Amount": "2500000", "Period": "2023", "Notes": "All company assets"},
    {"Category": "Liabilities", "Item": "Total Liabilities", "Amount": "1200000", "Period": "2023", "Notes": "All debt and obligations"}
  ]
}

IMPORTANT INSTRUCTIONS:
1. Extract ALL financial figures from the document with their exact amounts
2. Categorize each item appropriately (Revenue, Expense, Profit, Assets, Liabilities, etc.)
3. Identify the time period for each item if available
4. Include context or description for each financial figure
5. Format numbers as plain numbers without currency symbols or commas
6. Extract any numerical values that represent financial data, even if the document structure is unclear

EXTRACT ONLY THE JSON:"""
    
    # Format the prompt with the document content
    formatted_prompt = prompt_template.format(
        file_name=file_name,
        text=combined_text[:10000]  # Limit text to avoid token limits
    )
    
    # Call Ollama with increased timeout for complex extraction
    try:
        response = await call_ollama(formatted_prompt, config.OLLAMA_MODEL, timeout=60, max_retries=3, retry_delay=5)
        
        # Extract the JSON part
        if "{" in response:
            json_start = response.find("{")
            json_end = response.rfind("}") + 1
            if json_start >= 0 and json_end > json_start:
                try:
                    json_str = response[json_start:json_end]
                    extracted_data = json.loads(json_str)
                    
                    # Validate the extracted data structure
                    if "columns" in extracted_data and "data" in extracted_data and len(extracted_data["data"]) > 0:
                        logger.info(f"Successfully extracted {len(extracted_data['data'])} rows of data using Ollama agent")
                        return extracted_data
                except json.JSONDecodeError as e:
                    logger.warning(f"Invalid JSON in Ollama response: {str(e)}")
        
        # If we get here, either JSON parsing failed or the structure was invalid
        logger.warning("Failed to extract structured data from Ollama response")
        return None
        
    except Exception as e:
        logger.error(f"Error using Ollama for extraction: {str(e)}")
        return None

async def extract_metadata_using_ollama_agent(file_info, analysis_text):
    """
    Extract document metadata using Ollama as an agent-based approach
    
    Args:
        file_info: FileInfo object with document details
        analysis_text: Analysis text from document analysis
        
    Returns:
        Structured metadata as a dictionary
    """
    logger.info(f"Extracting metadata using Ollama agent approach for {file_info.name}")
    
    # Create a specialized prompt for metadata extraction
    prompt_template = """
You are a Financial Metadata Specialist working with an AI system. 
Your task is to extract comprehensive metadata from financial document analysis.

DOCUMENT: {file_name}
DOCUMENT ANALYSIS:
{analysis_text}

Extract the following metadata fields and provide them in JSON format:
{
  "report_type": "The document type (e.g., Balance Sheet, Income Statement, Annual Report)",
  "report_period": "The specific time period covered (e.g., Q1 2023, FY 2024)",
  "client_name": "The company or organization name",
  "entity": "The department or entity that produced the document",
  "account_name": "The account or division name if applicable",
  "wallet_id": "Any account identifier present, or generate a plausible one",
  "description": "A detailed 1-2 sentence summary of the document's contents",
  "information_present": ["list", "of", "key", "financial", "metrics", "present"]
}

IMPORTANT REQUIREMENTS:
1. The report_period MUST include a specific time period like "Q1 2023", "FY 2024", "Jan-Mar 2024"
2. The description MUST be a detailed 1-2 sentence summary of what the document contains
3. The information_present MUST list at least 3-5 specific financial metrics or data points found in the document
4. If a field isn't explicitly mentioned, make a reasonable inference based on context
5. Be comprehensive and detailed in your extraction
6. Format the output as valid JSON only

EXTRACT ONLY THE JSON:"""
    
    # Format the prompt with document information
    formatted_prompt = prompt_template.format(
        file_name=file_info.name,
        analysis_text=analysis_text
    )
    
    # Call Ollama with increased timeout for complex extraction
    try:
        response = await call_ollama(formatted_prompt, config.OLLAMA_MODEL, timeout=45, max_retries=2, retry_delay=3)
        
        # Extract the JSON part
        if "{" in response:
            json_start = response.find("{")
            json_end = response.rfind("}") + 1
            if json_start >= 0 and json_end > json_start:
                try:
                    json_str = response[json_start:json_end]
                    extracted_metadata = json.loads(json_str)
                    
                    # Validate and ensure all required fields are present
                    required_fields = ["report_type", "report_period", "client_name", "description", "information_present"]
                    for field in required_fields:
                        if field not in extracted_metadata or not extracted_metadata[field]:
                            if field == "report_period":
                                # Infer from filename or use current date
                                filename = file_info.name.lower()
                                current_year = datetime.now().year
                                current_quarter = (datetime.now().month - 1) // 3 + 1
                                if any(f"q{i}" in filename for i in range(1, 5)):
                                    for i in range(1, 5):
                                        if f"q{i}" in filename:
                                            quarter = i
                                            year_match = re.search(r'20\d{2}', filename)
                                            year = year_match.group(0) if year_match else str(current_year)
                                            extracted_metadata["report_period"] = f"Q{quarter} {year}"
                                            break
                                else:
                                    extracted_metadata["report_period"] = f"Q{current_quarter} {current_year}"
                            
                            elif field == "client_name":
                                extracted_metadata["client_name"] = "Unknown Entity"
                            
                            elif field == "description":
                                # Create description based on report type
                                report_type = extracted_metadata.get("report_type", "Unknown").lower()
                                if "balance sheet" in report_type:
                                    extracted_metadata["description"] = "Statement of financial position showing assets, liabilities, and equity balances."
                                elif "income statement" in report_type or "profit" in report_type:
                                    extracted_metadata["description"] = "Statement of profit and loss detailing revenue, expenses, and net income."
                                elif "cash flow" in report_type:
                                    extracted_metadata["description"] = "Statement of cash flows showing operating, investing, and financing activities."
                                elif "annual report" in report_type:
                                    extracted_metadata["description"] = "Comprehensive annual report with financial statements and company performance analysis."
                                else:
                                    extracted_metadata["description"] = f"Financial document containing {report_type} information and analysis."
                            
                            elif field == "information_present":
                                # Default information based on report type
                                report_type = extracted_metadata.get("report_type", "Unknown").lower()
                                if "balance sheet" in report_type:
                                    extracted_metadata["information_present"] = ["assets", "liabilities", "equity", "total assets", "current assets"]
                                elif "income statement" in report_type or "profit" in report_type:
                                    extracted_metadata["information_present"] = ["revenue", "expenses", "gross profit", "net income", "taxes"]
                                elif "cash flow" in report_type:
                                    extracted_metadata["information_present"] = ["operating activities", "investing activities", "financing activities", "net cash flow", "cash balance"]
                                else:
                                    extracted_metadata["information_present"] = ["financial data", "metrics", "analysis", "performance indicators", "summaries"]
                    
                    # Generate default wallet_id if missing
                    if "wallet_id" not in extracted_metadata or not extracted_metadata["wallet_id"]:
                        file_id = abs(hash(file_info.path)) % 10000
                        extracted_metadata["wallet_id"] = f"WLT_{file_id:04d}"
                    
                    # Add entity if missing
                    if "entity" not in extracted_metadata or not extracted_metadata["entity"]:
                        extracted_metadata["entity"] = "Finance Department"
                    
                    # Add account_name if missing
                    if "account_name" not in extracted_metadata or not extracted_metadata["account_name"]:
                        extracted_metadata["account_name"] = "Primary Account"
                    
                    # Ensure at least 3 items in information_present
                    if len(extracted_metadata.get("information_present", [])) < 3:
                        extracted_metadata["information_present"] = list(extracted_metadata.get("information_present", []))
                        extracted_metadata["information_present"].extend(["financial data", "metrics", "analysis"])
                        # Remove duplicates while preserving order
                        extracted_metadata["information_present"] = list(dict.fromkeys(extracted_metadata["information_present"]))
                    
                    logger.info(f"Successfully extracted metadata using Ollama agent for {file_info.name}")
                    return extracted_metadata
                    
                except json.JSONDecodeError as e:
                    logger.warning(f"Invalid JSON in Ollama response: {str(e)}")
        
        # If we get here, extraction failed
        logger.warning(f"Failed to extract metadata using Ollama agent for {file_info.name}")
        return None
        
    except Exception as e:
        logger.error(f"Error using Ollama for metadata extraction: {str(e)}")
        return None